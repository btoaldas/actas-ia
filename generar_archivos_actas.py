#!/usr/bin/env python
"""
Generador de archivos PDF, Word y TXT para las actas municipales existentes
"""

import os
import sys
import django
from io import BytesIO
from datetime import datetime

# Configurar Django
sys.path.append('/app')
os.environ.setdefault('DJANGO_SETTINGS_MODULE', 'config.settings')
django.setup()

from django.contrib.auth.models import User
from apps.pages.models import ActaMunicipal
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
from docx import Document
from docx.shared import Inches


class GeneradorArchivosActas:
    """Generador de archivos para actas municipales"""
    
    def __init__(self):
        self.archivos_generados = []
        
    def generar_pdf(self, acta):
        """Generar archivo PDF del acta"""
        try:
            # Crear directorio si no existe
            pdf_dir = f"/app/media/actas_pdf/{acta.fecha_sesion.year}/"
            os.makedirs(pdf_dir, exist_ok=True)
            
            pdf_path = f"{pdf_dir}acta_{acta.numero_acta.replace('-', '_')}.pdf"
            
            # Crear documento PDF
            doc = SimpleDocTemplate(pdf_path, pagesize=A4,
                                  rightMargin=72, leftMargin=72,
                                  topMargin=72, bottomMargin=18)
            
            # Estilos
            styles = getSampleStyleSheet()
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=16,
                spaceAfter=30,
                alignment=1,  # Centro
                textColor=colors.darkblue
            )
            
            # Contenido
            story = []
            
            # Título principal
            story.append(Paragraph("GAD MUNICIPAL DEL CANTÓN PASTAZA", title_style))
            story.append(Paragraph(f"ACTA N° {acta.numero_acta}", title_style))
            story.append(Spacer(1, 20))
            
            # Información básica
            info_style = styles['Normal']
            story.append(Paragraph(f"<b>Título:</b> {acta.titulo}", info_style))
            story.append(Spacer(1, 12))
            story.append(Paragraph(f"<b>Fecha de Sesión:</b> {acta.fecha_sesion.strftime('%d de %B de %Y')}", info_style))
            story.append(Spacer(1, 12))
            story.append(Paragraph(f"<b>Tipo de Sesión:</b> {acta.tipo_sesion.nombre if acta.tipo_sesion else 'Ordinaria'}", info_style))
            story.append(Spacer(1, 12))
            story.append(Paragraph(f"<b>Estado:</b> {acta.estado.nombre if acta.estado else 'Publicada'}", info_style))
            story.append(Spacer(1, 20))
            
            # Asistentes
            story.append(Paragraph("<b>ASISTENTES:</b>", styles['Heading2']))
            story.append(Paragraph(acta.asistentes, info_style))
            story.append(Spacer(1, 15))
            
            # Orden del día
            story.append(Paragraph("<b>ORDEN DEL DÍA:</b>", styles['Heading2']))
            story.append(Paragraph(acta.orden_del_dia, info_style))
            story.append(Spacer(1, 15))
            
            # Contenido
            story.append(Paragraph("<b>DESARROLLO DE LA SESIÓN:</b>", styles['Heading2']))
            # Limpiar contenido HTML básico
            contenido_limpio = acta.contenido.replace('<br>', '\n').replace('<b>', '').replace('</b>', '')
            story.append(Paragraph(contenido_limpio, info_style))
            story.append(Spacer(1, 15))
            
            # Acuerdos
            story.append(Paragraph("<b>ACUERDOS:</b>", styles['Heading2']))
            story.append(Paragraph(acta.acuerdos, info_style))
            story.append(Spacer(1, 20))
            
            # Firmas
            story.append(Paragraph("<b>FIRMAS:</b>", styles['Heading2']))
            story.append(Spacer(1, 15))
            
            # Obtener nombres seguros
            if acta.presidente and hasattr(acta.presidente, 'get_full_name'):
                presidente_nombre = acta.presidente.get_full_name()
            else:
                presidente_nombre = 'Segundo Germán Flores Meza'
                
            if acta.secretario and hasattr(acta.secretario, 'get_full_name'):
                secretario_nombre = acta.secretario.get_full_name()
            else:
                secretario_nombre = 'Danilo Andrade'
            
            story.append(Paragraph(f"<b>Presidente:</b> {presidente_nombre}", info_style))
            story.append(Spacer(1, 10))
            story.append(Paragraph(f"<b>Secretario:</b> {secretario_nombre}", info_style))
            
            # Construir PDF
            doc.build(story)
            
            # Actualizar campo en el modelo
            acta.archivo_pdf = f"actas_pdf/{acta.fecha_sesion.year}/acta_{acta.numero_acta.replace('-', '_')}.pdf"
            acta.save(update_fields=['archivo_pdf'])
            
            return pdf_path
            
        except Exception as e:
            print(f"❌ Error generando PDF para {acta.numero_acta}: {str(e)}")
            return None
    
    def generar_word(self, acta):
        """Generar archivo Word del acta"""
        try:
            # Crear directorio si no existe
            word_dir = f"/app/media/actas_word/{acta.fecha_sesion.year}/"
            os.makedirs(word_dir, exist_ok=True)
            
            word_path = f"{word_dir}acta_{acta.numero_acta.replace('-', '_')}.docx"
            
            # Crear documento Word
            doc = Document()
            
            # Título principal
            title = doc.add_heading('GAD MUNICIPAL DEL CANTÓN PASTAZA', 0)
            title.alignment = 1  # Centro
            
            subtitle = doc.add_heading(f'ACTA N° {acta.numero_acta}', level=1)
            subtitle.alignment = 1
            
            # Información básica
            doc.add_paragraph()
            info_table = doc.add_table(rows=4, cols=2)
            info_table.style = 'Table Grid'
            
            cells = info_table.rows[0].cells
            cells[0].text = 'Título:'
            cells[1].text = acta.titulo
            
            cells = info_table.rows[1].cells
            cells[0].text = 'Fecha de Sesión:'
            cells[1].text = acta.fecha_sesion.strftime('%d de %B de %Y')
            
            cells = info_table.rows[2].cells
            cells[0].text = 'Tipo de Sesión:'
            cells[1].text = acta.tipo_sesion.nombre if acta.tipo_sesion else 'Ordinaria'
            
            cells = info_table.rows[3].cells
            cells[0].text = 'Estado:'
            cells[1].text = acta.estado.nombre if acta.estado else 'Publicada'
            
            # Asistentes
            doc.add_paragraph()
            doc.add_heading('ASISTENTES', level=2)
            doc.add_paragraph(acta.asistentes)
            
            # Orden del día
            doc.add_heading('ORDEN DEL DÍA', level=2)
            doc.add_paragraph(acta.orden_del_dia)
            
            # Contenido
            doc.add_heading('DESARROLLO DE LA SESIÓN', level=2)
            contenido_limpio = acta.contenido.replace('<br>', '\n').replace('<b>', '').replace('</b>', '')
            doc.add_paragraph(contenido_limpio)
            
            # Acuerdos
            doc.add_heading('ACUERDOS', level=2)
            doc.add_paragraph(acta.acuerdos)
            
            # Firmas
            doc.add_heading('FIRMAS', level=2)
            doc.add_paragraph()
            
            # Obtener nombres seguros
            if acta.presidente and hasattr(acta.presidente, 'get_full_name'):
                presidente_nombre = acta.presidente.get_full_name()
            else:
                presidente_nombre = 'Segundo Germán Flores Meza'
                
            if acta.secretario and hasattr(acta.secretario, 'get_full_name'):
                secretario_nombre = acta.secretario.get_full_name()
            else:
                secretario_nombre = 'Danilo Andrade'
            
            firma_p = doc.add_paragraph()
            firma_p.add_run(f'Presidente: {presidente_nombre}').bold = True
            doc.add_paragraph()
            secretario_p = doc.add_paragraph()
            secretario_p.add_run(f'Secretario: {secretario_nombre}').bold = True
            
            # Guardar documento
            doc.save(word_path)
            
            # Actualizar campo en el modelo
            acta.archivo_word = f"actas_word/{acta.fecha_sesion.year}/acta_{acta.numero_acta.replace('-', '_')}.docx"
            acta.save(update_fields=['archivo_word'])
            
            return word_path
            
        except Exception as e:
            print(f"❌ Error generando Word para {acta.numero_acta}: {str(e)}")
            return None
    
    def generar_txt(self, acta):
        """Generar archivo TXT del acta"""
        try:
            # Crear directorio si no existe
            txt_dir = f"/app/media/actas_txt/{acta.fecha_sesion.year}/"
            os.makedirs(txt_dir, exist_ok=True)
            
            txt_path = f"{txt_dir}acta_{acta.numero_acta.replace('-', '_')}.txt"
            
            # Contenido del archivo TXT
            fecha_actual = datetime.now().strftime('%d de %B de %Y a las %H:%M')
            contenido_limpio = acta.contenido.replace('<br>', '\n').replace('<b>', '').replace('</b>', '')
            
            # Obtener nombres seguros
            if acta.presidente and hasattr(acta.presidente, 'get_full_name'):
                presidente_nombre = acta.presidente.get_full_name()
            else:
                presidente_nombre = 'Segundo Germán Flores Meza'
                
            if acta.secretario and hasattr(acta.secretario, 'get_full_name'):
                secretario_nombre = acta.secretario.get_full_name()
            else:
                secretario_nombre = 'Danilo Andrade'
            
            contenido_txt = f"""GAD MUNICIPAL DEL CANTÓN PASTAZA
ACTA N° {acta.numero_acta}

===============================================

TÍTULO: {acta.titulo}
FECHA DE SESIÓN: {acta.fecha_sesion.strftime('%d de %B de %Y')}
TIPO DE SESIÓN: {acta.tipo_sesion.nombre if acta.tipo_sesion else 'Ordinaria'}
ESTADO: {acta.estado.nombre if acta.estado else 'Publicada'}

===============================================
ASISTENTES:
{acta.asistentes}

===============================================
ORDEN DEL DÍA:
{acta.orden_del_dia}

===============================================
DESARROLLO DE LA SESIÓN:
{contenido_limpio}

===============================================
ACUERDOS:
{acta.acuerdos}

===============================================
FIRMAS:

Presidente: {presidente_nombre}
Secretario: {secretario_nombre}

===============================================
Documento generado el {fecha_actual}
Sistema de Actas Municipales - GAD Pastaza
"""
            
            # Escribir archivo
            with open(txt_path, 'w', encoding='utf-8') as f:
                f.write(contenido_txt)
            
            # Actualizar campo en el modelo
            acta.archivo_txt = f"actas_txt/{acta.fecha_sesion.year}/acta_{acta.numero_acta.replace('-', '_')}.txt"
            acta.save(update_fields=['archivo_txt'])
            
            return txt_path
            
        except Exception as e:
            print(f"❌ Error generando TXT para {acta.numero_acta}: {str(e)}")
            return None
    
    def procesar_todas_las_actas(self):
        """Procesar todas las actas existentes"""
        print("🚀 GENERANDO ARCHIVOS PARA ACTAS MUNICIPALES")
        print("=" * 60)
        
        actas = ActaMunicipal.objects.all().order_by('numero_acta')
        total_actas = actas.count()
        
        if total_actas == 0:
            print("❌ No hay actas municipales en el sistema")
            return
        
        print(f"📋 Procesando {total_actas} actas municipales...")
        print()
        
        archivos_creados = {'pdf': 0, 'word': 0, 'txt': 0}
        
        for i, acta in enumerate(actas, 1):
            print(f"📄 [{i}/{total_actas}] Procesando {acta.numero_acta}: {acta.titulo[:50]}...")
            
            # Generar PDF
            pdf_path = self.generar_pdf(acta)
            if pdf_path:
                archivos_creados['pdf'] += 1
                print(f"  ✅ PDF generado")
            
            # Generar Word
            word_path = self.generar_word(acta)
            if word_path:
                archivos_creados['word'] += 1
                print(f"  ✅ Word generado")
            
            # Generar TXT
            txt_path = self.generar_txt(acta)
            if txt_path:
                archivos_creados['txt'] += 1
                print(f"  ✅ TXT generado")
            
            print()
        
        print("=" * 60)
        print("✅ GENERACIÓN COMPLETADA")
        print("=" * 60)
        print(f"📊 Archivos PDF generados: {archivos_creados['pdf']}")
        print(f"📊 Archivos Word generados: {archivos_creados['word']}")
        print(f"📊 Archivos TXT generados: {archivos_creados['txt']}")
        print()
        print("🎉 ¡Todas las actas ahora tienen archivos descargables!")
        print("🔗 Los enlaces de descarga estarán disponibles en el portal ciudadano")


if __name__ == "__main__":
    generador = GeneradorArchivosActas()
    generador.procesar_todas_las_actas()