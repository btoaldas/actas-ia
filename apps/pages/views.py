from django.shortcuts import render, redirect, get_object_or_404
from apps.pages.forms import LoginForm, RegistrationForm, UserPasswordResetForm, UserSetPasswordForm, UserPasswordChangeForm
from django.contrib.auth import logout
from django.contrib.auth import views as auth_views
from django.db.models import Q, Count, Avg
from django.core.paginator import Paginator
from django.http import JsonResponse, HttpResponse, Http404
from django.views.decorators.csrf import csrf_exempt, csrf_protect
from django.views.decorators.http import require_http_methods
from django.utils.decorators import method_decorator
from django.contrib import messages
from django.views.generic import ListView, DetailView
from django.utils import timezone
from .models import ActaMunicipal, TipoSesion, EstadoActa, VisualizacionActa, DescargaActa
from helpers.util import normalizar_busqueda, crear_filtros_busqueda_multiple
import json
import os
import tempfile
import logging
from datetime import datetime, timedelta

logger = logging.getLogger(__name__)

# Create your views here.

# Authentication
def register(request):
  if request.method == 'POST':
    form = RegistrationForm(request.POST)
    if form.is_valid():
      form.save()
      print('Account created successfully!')
      return redirect('/accounts/login/')
    else:
      print("Registration failed!")
  else:
    form = RegistrationForm()
  
  context = {'form': form}
  return render(request, 'accounts/register.html', context)
  
def register_v1(request):
  if request.method == 'POST':
    form = RegistrationForm(request.POST)
    if form.is_valid():
      form.save()
      print('Account created successfully!')
      return redirect('/accounts/login/')
    else:
      print("Registration failed!")
  else:
    form = RegistrationForm()
  
  context = {'form': form}
  return render(request, 'pages/examples/register.html', context)

def register_v2(request):
  if request.method == 'POST':
    form = RegistrationForm(request.POST)
    if form.is_valid():
      form.save()
      print('Account created successfully!')
      return redirect('/accounts/login/')
    else:
      print("Registration failed!")
  else:
    form = RegistrationForm()
  
  context = {'form': form}
  return render(request, 'pages/examples/register-v2.html', context)

class UserLoginView(auth_views.LoginView):
  template_name = 'accounts/login.html'
  form_class = LoginForm
  success_url = '/'

class UserLoginViewV1(auth_views.LoginView):
  template_name = 'pages/examples/login.html'
  form_class = LoginForm
  success_url = '/'

class UserLoginViewV2(auth_views.LoginView):
  template_name = 'pages/examples/login-v2.html'
  form_class = LoginForm
  success_url = '/'

class UserPasswordResetView(auth_views.PasswordResetView):
  template_name = 'accounts/forgot-password.html'
  form_class = UserPasswordResetForm

class UserPasswordResetViewV1(auth_views.PasswordResetView):
  template_name = 'pages/examples/forgot-password.html'
  form_class = UserPasswordResetForm

class UserPasswordResetViewV2(auth_views.PasswordResetView):
  template_name = 'pages/examples/forgot-password-v2.html'
  form_class = UserPasswordResetForm

class UserPasswordResetConfirmView(auth_views.PasswordResetConfirmView):
  template_name = 'accounts/recover-password.html'
  form_class = UserSetPasswordForm

class UserPasswordChangeView(auth_views.PasswordChangeView):
  template_name = 'accounts/password_change.html'
  form_class = UserPasswordChangeForm

class UserPasswordChangeViewV1(auth_views.PasswordChangeView):
  template_name = 'pages/examples/recover-password.html'
  form_class = UserPasswordChangeForm

class UserPasswordChangeViewV2(auth_views.PasswordChangeView):
  template_name = 'pages/examples/recover-password-v2.html'
  form_class = UserPasswordChangeForm

def user_logout_view(request):
  logout(request)
  return redirect('/accounts/login/')


# pages
def index(request):
  # Si el usuario NO está autenticado, mostrar página de inicio pública
  if not request.user.is_authenticated:
    return inicio_publico(request)
  
  # Si el usuario SÍ está autenticado, mostrar el dashboard normal
  context = {
    'parent': 'dashboard',
    'segment': 'dashboardv1'
  }
  return render(request, 'pages/index.html', context)

def inicio_publico(request):
  """Página de inicio para usuarios no autenticados"""
  from .models import ActaMunicipal
  from django.db.models import Count
  from django.utils import timezone
  
  # Estadísticas públicas (solo números generales, sin detalles)
  stats_publicas = {
    'total_actas_publicas': ActaMunicipal.objects.filter(
      estado__nombre='publicada',
      acceso='publico'
    ).count(),
    'total_sesiones_ano': ActaMunicipal.objects.filter(
      estado__nombre='publicada',
      acceso='publico',
      fecha_sesion__year=timezone.now().year
    ).count(),
  }
  
  # Últimas actas públicas (máximo 3)
  ultimas_actas = ActaMunicipal.objects.filter(
    estado__nombre='publicada',
    acceso='publico'
  ).order_by('-fecha_publicacion')[:3]
  
  context = {
    'stats_publicas': stats_publicas,
    'ultimas_actas': ultimas_actas,
    'mostrar_informacion_publica': True,
  }
  return render(request, 'pages/inicio_publico_completo.html', context)

def index2(request):
  # Datos simulados para el dashboard de actas municipales
  from datetime import datetime, timedelta
  import json
  
  # Estadísticas generales del sistema
  estadisticas = {
    'actas_totales': 1247,
    'actas_mes_actual': 89,
    'sesiones_pendientes': 5,
    'documentos_procesados': 2891,
    'precision_ia': 96.7,
    'tiempo_medio_transcripcion': 12.3,
    'usuarios_activos': 23,
    'almacenamiento_usado': 78.4
  }
  
  # Datos para gráficos (últimos 12 meses)
  meses = ['Ene', 'Feb', 'Mar', 'Abr', 'May', 'Jun', 
           'Jul', 'Ago', 'Sep', 'Oct', 'Nov', 'Dic']
  actas_por_mes = [67, 84, 91, 102, 78, 95, 112, 89, 76, 103, 87, 89]
  transcripciones_ia = [45, 62, 78, 89, 67, 84, 97, 76, 68, 91, 79, 82]
  
  # Actividad reciente
  actividad_reciente = [
    {
      'tipo': 'transcripcion',
      'descripcion': 'Sesión Ordinaria - Consejo Municipal',
      'fecha': datetime.now() - timedelta(hours=2),
      'usuario': 'Ana López',
      'estado': 'completada'
    },
    {
      'tipo': 'revision',
      'descripcion': 'Acta Extraordinaria - Presupuesto 2025',
      'fecha': datetime.now() - timedelta(hours=5),
      'usuario': 'Juan Pérez',
      'estado': 'pendiente'
    },
    {
      'tipo': 'publicacion',
      'descripcion': 'Acta Ordinaria - Urbanismo',
      'fecha': datetime.now() - timedelta(hours=8),
      'usuario': 'María García',
      'estado': 'publicada'
    },
    {
      'tipo': 'procesamiento',
      'descripción': 'Análisis IA - Sesión de Hacienda',
      'fecha': datetime.now() - timedelta(days=1),
      'usuario': 'Sistema IA',
      'estado': 'procesando'
    }
  ]
  
  # Próximas sesiones programadas
  proximas_sesiones = [
    {
      'titulo': 'Consejo Municipal Ordinario',
      'fecha': datetime.now() + timedelta(days=3),
      'hora': '10:00',
      'sala': 'Salón de Plenos',
      'tipo': 'ordinaria'
    },
    {
      'titulo': 'Comisión de Urbanismo',
      'fecha': datetime.now() + timedelta(days=7),
      'hora': '16:30',
      'sala': 'Sala de Comisiones',
      'tipo': 'comision'
    },
    {
      'titulo': 'Sesión Extraordinaria - Presupuestos',
      'fecha': datetime.now() + timedelta(days=12),
      'hora': '09:00',
      'sala': 'Salón de Plenos',
      'tipo': 'extraordinaria'
    }
  ]
  
  # Estadísticas de precisión IA por categoría
  precision_categorias = {
    'Resoluciones': 98.2,
    'Debates': 94.7,
    'Votaciones': 99.1,
    'Intervenciones': 95.3,
    'Documentos técnicos': 97.8
  }
  
  context = {
    'parent': 'dashboard',
    'segment': 'dashboardv2',
    'estadisticas': estadisticas,
    'meses_json': json.dumps(meses),
    'actas_por_mes_json': json.dumps(actas_por_mes),
    'transcripciones_ia_json': json.dumps(transcripciones_ia),
    'actividad_reciente': actividad_reciente,
    'proximas_sesiones': proximas_sesiones,
    'precision_categorias': precision_categorias
  }
  return render(request, 'pages/index2.html', context)

def index3(request):
  # Dashboard avanzado enfocado en análisis y transparencia
  from datetime import datetime, timedelta
  import json
  
  # Análisis de transparencia y cumplimiento normativo
  transparencia_metrics = {
    'actas_publicadas_24h': 87.3,  # % publicadas en 24h
    'cumplimiento_plazos': 94.1,   # % cumplimiento plazos legales
    'accesibilidad_web': 98.5,     # % accesibilidad WCAG
    'tiempo_medio_publicacion': 18.7,  # horas promedio
    'documentos_traducidos': 76.2,  # % traducidos a idiomas cooficiales
    'quejas_ciudadanas': 3,         # quejas por transparencia
    'solicitudes_info': 142,       # solicitudes de información
    'descargas_mensuales': 8947     # descargas de actas
  }
  
  # Análisis de contenido por IA
  analisis_contenido = {
    'temas_frecuentes': [
      {'tema': 'Presupuestos y Hacienda', 'frecuencia': 34, 'tendencia': '+12%'},
      {'tema': 'Urbanismo y Obras', 'frecuencia': 28, 'tendencia': '+8%'},
      {'tema': 'Servicios Sociales', 'frecuencia': 19, 'tendencia': '-3%'},
      {'tema': 'Medio Ambiente', 'frecuencia': 15, 'tendencia': '+25%'},
      {'tema': 'Cultura y Deportes', 'frecuencia': 11, 'tendencia': '+5%'},
      {'tema': 'Seguridad Ciudadana', 'frecuencia': 9, 'tendencia': '-7%'}
    ],
    'sentimientos_detectados': {
      'positivo': 58.3,
      'neutral': 35.2,
      'negativo': 6.5
    },
    'palabras_clave_mes': ['presupuesto', 'moción', 'aprobación', 'debate', 'votación']
  }
  
  # Métricas de calidad del sistema
  calidad_sistema = {
    'precision_transcripcion': 96.8,
    'precision_clasificacion': 94.2,
    'precision_extraccion_datos': 97.1,
    'tiempo_procesamiento': 12.3,
    'disponibilidad_sistema': 99.7,
    'satisfaccion_usuario': 4.6  # sobre 5
  }
  
  # Datos para gráficos de tendencias (últimos 6 meses)
  meses_trend = ['Abr', 'May', 'Jun', 'Jul', 'Ago', 'Sep']
  transparencia_trend = [82.1, 85.4, 88.2, 90.1, 92.3, 94.1]
  publicaciones_trend = [78, 85, 91, 89, 94, 87]
  accesos_ciudadanos = [1840, 2120, 2350, 2180, 2560, 2890]
  
  # Alertas y notificaciones del sistema
  alertas_sistema = [
    {
      'tipo': 'warning',
      'titulo': 'Plazo de publicación próximo',
      'descripcion': 'Acta del 4 de septiembre debe publicarse antes de las 18:00',
      'tiempo': '2 horas restantes',
      'urgencia': 'alta'
    },
    {
      'tipo': 'info',
      'titulo': 'Actualización de modelo IA',
      'descripcion': 'Nueva versión del modelo de transcripción disponible',
      'tiempo': 'Hace 1 día',
      'urgencia': 'media'
    },
    {
      'tipo': 'success',
      'titulo': 'Respaldo completado',
      'descripcion': 'Respaldo automático de documentos finalizado correctamente',
      'tiempo': 'Hace 3 horas',
      'urgencia': 'baja'
    },
    {
      'tipo': 'danger',
      'titulo': 'Error en transcripción',
      'descripcion': 'Sesión del 5 de septiembre requiere revisión manual',
      'tiempo': 'Hace 30 minutos',
      'urgencia': 'crítica'
    }
  ]
  
  # Estadísticas de uso ciudadano
  uso_ciudadano = {
    'visitantes_unicos': 12847,
    'busquedas_realizadas': 3962,
    'documentos_descargados': 8947,
    'tiempo_promedio_sesion': '4m 32s',
    'dispositivos_moviles': 68.3,  # %
    'horario_pico': '10:00-12:00'
  }
  
  # Cumplimiento normativo
  cumplimiento_normativo = [
    {'norma': 'Ley 19/2013 Transparencia', 'cumplimiento': 97.2, 'estado': 'excelente'},
    {'norma': 'Ley 39/2015 Procedimiento', 'cumplimiento': 94.8, 'estado': 'muy_bueno'},
    {'norma': 'RGPD Protección Datos', 'cumplimiento': 98.1, 'estado': 'excelente'},
    {'norma': 'Real Decreto 203/2021', 'cumplimiento': 91.3, 'estado': 'bueno'},
    {'norma': 'UNE 139803 Accesibilidad', 'cumplimiento': 96.7, 'estado': 'muy_bueno'}
  ]
  
  context = {
    'parent': 'dashboard',
    'segment': 'dashboardv3',
    'transparencia_metrics': transparencia_metrics,
    'analisis_contenido': analisis_contenido,
    'calidad_sistema': calidad_sistema,
    'meses_trend_json': json.dumps(meses_trend),
    'transparencia_trend_json': json.dumps(transparencia_trend),
    'publicaciones_trend_json': json.dumps(publicaciones_trend),
    'accesos_ciudadanos_json': json.dumps(accesos_ciudadanos),
    'alertas_sistema': alertas_sistema,
    'uso_ciudadano': uso_ciudadano,
    'cumplimiento_normativo': cumplimiento_normativo
  }
  return render(request, 'pages/index3.html', context)

def widgets(request):
  context = {
    'parent': '',
    'segment': 'widgets'
  }
  return render(request, 'pages/widgets.html', context)

# EXAMPLES

def examples_calendar(request):
  context = {
    'parent': '',
    'segment': 'calendar'
  }
  return render(request, 'pages/calendar.html', context)

def examples_gallery(request):
  context = {
    'parent': '',
    'segment': 'gallery'
  }
  return render(request, 'pages/gallery.html', context)

def examples_kanban(request):
  context = {
    'parent': '',
    'segment': 'kanban_board'
  }
  return render(request, 'pages/kanban.html', context)

# Mailbox

def mailbox_inbox(request):
  context = {
    'parent': 'mailbox',
    'segment': 'inbox'
  }
  return render(request, 'pages/mailbox/mailbox.html', context)

def mailbox_compose(request):
  context = {
    'parent': 'mailbox',
    'segment': 'compose'
  }
  return render(request, 'pages/mailbox/compose.html', context)

def mailbox_read_mail(request):
  context = {
    'parent': 'mailbox',
    'segment': 'read_mail'
  }
  return render(request, 'pages/mailbox/read-mail.html', context)

# Pages -- Examples

def examples_invoice(request):
  context = {
    'parent': 'pages',
    'segment': 'invoice'
  }
  return render(request, 'pages/examples/invoice.html', context)

def invoice_print(request):
  context = {
    'parent': 'pages',
    'segment': 'invoice_print'
  }
  return render(request, 'pages/examples/invoice-print.html', context)

def examples_profile(request):
  context = {
    'parent': 'pages',
    'segment': 'profile'
  }
  return render(request, 'pages/examples/profile.html', context)

def examples_e_commerce(request):
  context = {
    'parent': 'pages',
    'segment': 'ecommerce'
  }
  return render(request, 'pages/examples/e-commerce.html', context)

def examples_projects(request):
  context = {
    'parent': 'pages',
    'segment': 'projects'
  }
  return render(request, 'pages/examples/projects.html', context)

def examples_project_add(request):
  context = {
    'parent': 'pages',
    'segment': 'project_add'
  }
  return render(request, 'pages/examples/project-add.html', context)

def examples_project_edit(request):
  context = {
    'parent': 'pages',
    'segment': 'project_edit'
  }
  return render(request, 'pages/examples/project-edit.html', context)

def examples_project_detail(request):
  context = {
    'parent': 'pages',
    'segment': 'project_detail'
  }
  return render(request, 'pages/examples/project-detail.html', context)

def examples_contacts(request):
  context = {
    'parent': 'pages',
    'segment': 'contacts'
  }
  return render(request, 'pages/examples/contacts.html', context)

def examples_faq(request):
  context = {
    'parent': 'pages',
    'segment': 'faq'
  }
  return render(request, 'pages/examples/faq.html', context)

def examples_contact_us(request):
  context = {
    'parent': 'pages',
    'segment': 'contact_us'
  }
  return render(request, 'pages/examples/contact-us.html', context)

# Extra -- login & Registration v1
# def login_v1(request):
#   context = {
#     'parent': '',
#     'segment': ''
#   }
#   return render(request, 'pages/examples/login.html', context)

# def login_v2(request):
#   context = {
#     'parent': '',
#     'segment': ''
#   }
#   return render(request, 'pages/examples/login-v2.html', context)

# def registration_v1(request):
#   context = {
#     'parent': '',
#     'segment': ''
#   }
#   return render(request, 'pages/examples/register.html', context)

# def registration_v2(request):
#   context = {
#     'parent': '',
#     'segment': ''
#   }
#   return render(request, 'pages/examples/register-v2.html', context)

# def forgot_password_v1(request):
#   context = {
#     'parent': '',
#     'segment': ''
#   }
#   return render(request, 'pages/examples/forgot-password.html', context)

# def forgot_password_v2(request):
#   context = {
#     'parent': '',
#     'segment': ''
#   }
#   return render(request, 'pages/examples/forgot-password-v2.html', context)

# def recover_password_v1(request):
#   context = {
#     'parent': '',
#     'segment': ''
#   }
#   return render(request, 'pages/examples/recover-password.html', context)

# def recover_password_v2(request):
#   context = {
#     'parent': '',
#     'segment': ''
#   }
#   return render(request, 'pages/examples/recover-password-v2.html', context)

def lockscreen(request):
  context = {
    'parent': '',
    'segment': ''
  }
  return render(request, 'pages/examples/lockscreen.html', context)

def legacy_user_menu(request):
  context = {
    'parent': 'extra',
    'segment': 'legacy_user'
  }
  return render(request, 'pages/examples/legacy-user-menu.html', context)

def language_menu(request):
  context = {
    'parent': 'extra',
    'segment': 'legacy_menu'
  }
  return render(request, 'pages/examples/language-menu.html', context)

def error_404(request):
  context = {
    'parent': 'extra',
    'segment': 'error_404'
  }
  return render(request, 'pages/examples/404.html', context)

def error_500(request):
  context = {
    'parent': 'extra',
    'segment': 'error_500'
  }
  return render(request, 'pages/examples/500.html', context)

def pace(request):
  context = {
    'parent': 'extra',
    'segment': 'pace'
  }
  return render(request, 'pages/examples/pace.html', context)

def blank_page(request):
  context = {
    'parent': 'extra',
    'segment': 'blank_page'
  }
  return render(request, 'pages/examples/blank.html', context)

def starter_page(request):
  context = {
    'parent': 'extra',
    'segment': 'starter_page'
  }
  return render(request, 'pages/examples/starter.html', context)

# Search

def search_simple(request):
  context = {
    'parent': 'search',
    'segment': 'search_simple'
  }
  return render(request, 'pages/search/simple.html', context)

def search_enhanced(request):
  context = {
    'parent': 'search',
    'segment': 'search_enhanced'
  }
  return render(request, 'pages/search/enhanced.html', context)

def simple_results(request):
  context = {
    'parent': '',
    'segment': ''
  }
  return render(request, 'pages/search/simple-results.html', context)

def enhanced_results(request):
  context = {
    'parent': '',
    'segment': ''
  }
  return render(request, 'pages/search/enhanced-results.html', context)

# MISCELLANEOUS

def iframe(request):
  context = {
    'parent': '',
    'segment': ''
  }
  return render(request, 'pages/search/iframe.html', context)

# Charts

def chartjs(request):
  context = {
    'parent': 'charts',
    'segment': 'chartjs'
  }
  return render(request, 'pages/charts/chartjs.html', context)

def flot(request):
  context = {
    'parent': 'charts',
    'segment': 'flot'
  }
  return render(request, 'pages/charts/flot.html', context)

def inline(request):
  context = {
    'parent': 'charts',
    'segment': 'inline'
  }
  return render(request, 'pages/charts/inline.html', context)

def uplot(request):
  context = {
    'parent': 'charts',
    'segment': 'uplot'
  }
  return render(request, 'pages/charts/uplot.html', context)

def profile(request):
  context = {
    'parent': 'pages',
    'segment': 'profile'
  }
  return render(request, 'pages/examples/profile.html', context)

# Layout
def top_navigation(request):
  context = {
    'parent': 'layout',
    'segment': 'top_navigation'
  }
  return render(request, 'pages/layout/top-nav.html', context)

def top_nav_sidebar(request):
  context = {
    'parent': 'layout',
    'segment': 'top navigation with sidebar'
  }
  return render(request, 'pages/layout/top-nav-sidebar.html', context)

def boxed(request):
  context = {
    'parent': 'layout',
    'segment': 'boxed_layout'
  }
  return render(request, 'pages/layout/boxed.html', context)

def fixed_sidebar(request):
  context = {
    'parent': 'layout',
    'segment': 'fixed_layout'
  }
  return render(request, 'pages/layout/fixed-sidebar.html', context)

def fixed_sidebar_custom(request):
  context = {
    'parent': 'layout',
    'segment': 'layout_cuastom'
  }
  return render(request, 'pages/layout/fixed-sidebar-custom.html', context)

def fixed_topnav(request):
  context = {
    'parent': 'layout',
    'segment': 'fixed_topNav'
  }
  return render(request, 'pages/layout/fixed-topnav.html', context)

def fixed_navbar(request):
  context = {
    'parent': 'layout',
    'segment': 'fixed_navbar'
  }
  return render(request, 'pages/layout/fixed-navbar.html', context)

def fixed_footer(request):
  context = {
    'parent': 'layout',
    'segment': 'fixed_footer'
  }
  return render(request, 'pages/layout/fixed-footer.html', context)

def collapsed_sidebar(request):
  context = {
    'parent': 'layout',
    'segment': 'collapsed_sidebar'
  }
  return render(request, 'pages/layout/collapsed-sidebar.html', context)

# UI Elements

def ui_general(request):
  context = {
    'parent': 'ui',
    'segment': 'general'
  }
  return render(request, 'pages/UI/general.html', context)

def ui_icons(request):
  context = {
    'parent': 'ui',
    'segment': 'icons'
  }
  return render(request, 'pages/UI/icons.html', context)

def ui_buttons(request):
  context = {
    'parent': 'ui',
    'segment': 'buttons'
  }
  return render(request, 'pages/UI/buttons.html', context)

def ui_sliders(request):
  context = {
    'parent': 'ui',
    'segment': 'sliders'
  }
  return render(request, 'pages/UI/sliders.html', context)

def ui_modals_alerts(request):
  context = {
    'parent': 'ui',
    'segment': 'modals_alerts'
  }
  return render(request, 'pages/UI/modals.html', context)

def ui_navbar_tabs(request):
  context = {
    'parent': 'ui',
    'segment': 'navbar_tabs'
  }
  return render(request, 'pages/UI/navbar.html', context)

def ui_timeline(request):
  context = {
    'parent': 'ui',
    'segment': 'timeline'
  }
  return render(request, 'pages/UI/timeline.html', context)

def ui_ribbons(request):
  context = {
    'parent': 'ui',
    'segment': 'ribbons'
  }
  return render(request, 'pages/UI/ribbons.html', context)

# Forms

def form_general(request):
  context = {
    'parent': 'forms',
    'segment': 'form_general'
  }
  return render(request, 'pages/forms/general.html', context)

def form_advanced(request):
  context = {
    'parent': 'forms',
    'segment': 'advanced_form'
  }
  return render(request, 'pages/forms/advanced.html', context)

def form_editors(request):
  context = {
    'parent': 'forms',
    'segment': 'text_editors'
  }
  return render(request, 'pages/forms/editors.html', context)

def form_validation(request):
  context = {
    'parent': 'forms',
    'segment': 'validation'
  }
  return render(request, 'pages/forms/validation.html', context)

# Table

def table_simple(request):
  context = {
    'parent': 'tables',
    'segment': 'simple_table'
  }
  return render(request, 'pages/tables/simple.html', context)

def table_data(request):
  context = {
    'parent': 'tables',
    'segment': 'data_table'
  }
  return render(request, 'pages/tables/data.html', context)

def table_jsgrid(request):
  context = {
    'parent': 'tables',
    'segment': 'jsGrid'
  }
  return render(request, 'pages/tables/jsgrid.html', context)



def handler404(request, exception=None):
  return render(request, 'accounts/error-404.html')

def handler403(request, exception=None):
  return render(request, 'accounts/error-403.html')

def handler500(request, exception=None):
  return render(request, 'accounts/error-500.html')


# i18n
def i18n_view(request):
  context = {
    'parent': 'apps',
    'segment': 'i18n'
  }
  return render(request, 'pages/navigation/i18n.html', context)


# ============================================================================
# SISTEMA DE ACTAS MUNICIPALES - VISTAS DEMO
# ============================================================================

def actas_listado(request):
    """Vista demo para el listado de actas municipales"""
    context = {
        'parent': 'actas',
        'segment': 'listado',
        'page_title': 'Listado de Actas Municipales',
        'actas_demo': [
            {
                'id': 1,
                'numero': 'ACT-2024-001',
                'fecha': '2024-01-15',
                'tipo': 'Sesión Ordinaria',
                'estado': 'Publicada',
                'presidente': 'Ing. María Rodríguez',
                'participantes': 8,
                'estado_ia': 'Procesada',
                'confianza_ia': 95.2
            },
            {
                'id': 2,
                'numero': 'ACT-2024-002',
                'fecha': '2024-01-22',
                'tipo': 'Sesión Extraordinaria',
                'estado': 'En Revisión',
                'presidente': 'Ing. María Rodríguez',
                'participantes': 6,
                'estado_ia': 'En Proceso',
                'confianza_ia': 87.8
            },
            {
                'id': 3,
                'numero': 'ACT-2024-003',
                'fecha': '2024-02-05',
                'tipo': 'Sesión de Comisión',
                'estado': 'Borrador',
                'presidente': 'Dr. Carlos Mendoza',
                'participantes': 5,
                'estado_ia': 'Pendiente',
                'confianza_ia': 0
            }
        ]
    }
    return render(request, 'pages/actas/listado.html', context)

def actas_nueva(request):
    """Vista demo para crear nueva acta"""
    context = {
        'parent': 'actas',
        'segment': 'nueva',
        'page_title': 'Nueva Acta Municipal',
        'tipos_sesion': [
            'Sesión Ordinaria',
            'Sesión Extraordinaria', 
            'Sesión de Comisión',
            'Sesión de Emergencia'
        ],
        'funcionarios': [
            {'id': 1, 'nombre': 'Ing. María Rodríguez', 'cargo': 'Alcaldesa'},
            {'id': 2, 'nombre': 'Dr. Carlos Mendoza', 'cargo': 'Vicealcalde'},
            {'id': 3, 'nombre': 'Lcda. Ana García', 'cargo': 'Secretaria General'},
            {'id': 4, 'nombre': 'Ing. Luis Vargas', 'cargo': 'Director de Obras'},
        ]
    }
    return render(request, 'pages/actas/nueva.html', context)

def actas_transcripcion(request):
    """Vista demo para transcripción automática"""
    context = {
        'parent': 'ia_procesamiento',
        'segment': 'transcripcion',
        'page_title': 'Transcripción Automática de Audio',
        'archivos_demo': [
            {
                'nombre': 'sesion_ordinaria_20240115.mp3',
                'duracion': '2:45:30',
                'estado': 'Completado',
                'confianza': 96.5,
                'fecha_subida': '2024-01-15 14:30'
            },
            {
                'nombre': 'sesion_extraordinaria_20240122.wav',
                'duracion': '1:32:15',
                'estado': 'En Proceso',
                'confianza': 78.2,
                'fecha_subida': '2024-01-22 10:15'
            }
        ]
    }
    return render(request, 'pages/ia/transcripcion.html', context)

def actas_analisis_texto(request):
    """Vista demo para análisis de texto con IA"""
    context = {
        'parent': 'ia_procesamiento',
        'segment': 'analisis',
        'page_title': 'Análisis Inteligente de Texto',
        'analisis_demo': {
            'documento': 'ACT-2024-001',
            'resumen_automatico': 'La sesión trató principalmente sobre la aprobación del presupuesto municipal 2024, la autorización de obras de infraestructura vial y la implementación de nuevos programas sociales.',
            'temas_principales': [
                'Presupuesto Municipal 2024',
                'Infraestructura Vial',
                'Programas Sociales',
                'Ordenanzas Municipales'
            ],
            'participantes_detectados': [
                'Ing. María Rodríguez (Alcaldesa)',
                'Dr. Carlos Mendoza (Vicealcalde)',
                'Lcda. Ana García (Secretaria)',
                'Ing. Luis Vargas (Director de Obras)'
            ],
            'decisiones_tomadas': [
                'Aprobación del presupuesto por $2.5M',
                'Autorización obras Av. Principal',
                'Implementación programa alimentario'
            ],
            'confianza_general': 94.7
        }
    }
    return render(request, 'pages/ia/analisis_texto.html', context)

def actas_digitalizacion(request):
    """Vista demo para digitalización de documentos"""
    context = {
        'parent': 'digitalizacion',
        'segment': 'escaneo',
        'page_title': 'Digitalización de Documentos',
        'documentos_demo': [
            {
                'nombre': 'acta_manuscrita_1985.pdf',
                'paginas': 12,
                'calidad_ocr': 'Excelente',
                'estado': 'Procesado',
                'texto_extraido': 98.5,
                'fecha_proceso': '2024-01-10'
            },
            {
                'nombre': 'acta_manuscrita_1990.pdf', 
                'paginas': 8,
                'calidad_ocr': 'Buena',
                'estado': 'En Proceso',
                'texto_extraido': 85.2,
                'fecha_proceso': '2024-01-12'
            }
        ]
    }
    return render(request, 'pages/digitalizacion/escaneo.html', context)

def actas_publicacion(request):
    """Vista demo para publicación y transparencia"""
    context = {
        'parent': 'publicacion',
        'segment': 'transparencia',
        'page_title': 'Portal de Transparencia',
        'estadisticas': {
            'actas_publicadas': 245,
            'visitas_mes': 1250,
            'descargas_mes': 450,
            'ultima_actualizacion': '2024-01-22'
        },
        'actas_recientes': [
            {
                'numero': 'ACT-2024-001',
                'fecha': '2024-01-15',
                'titulo': 'Sesión Ordinaria - Aprobación Presupuesto 2024',
                'descargas': 85,
                'visitas': 320
            },
            {
                'numero': 'ACT-2024-002', 
                'fecha': '2024-01-22',
                'titulo': 'Sesión Extraordinaria - Infraestructura Vial',
                'descargas': 62,
                'visitas': 205
            }
        ]
    }
    return render(request, 'pages/publicacion/transparencia.html', context)

def actas_reportes(request):
    """Vista demo para reportes y análisis"""
    context = {
        'parent': 'reportes',
        'segment': 'metricas',
        'page_title': 'Métricas de Gestión Municipal',
        'metricas': {
            'eficiencia_proceso': 92.5,
            'tiempo_promedio_acta': '2.5 días',
            'precision_ia': 94.8,
            'satisfaccion_usuarios': 88.2
        },
        'tendencias': {
            'actas_por_mes': [12, 15, 18, 14, 16, 20, 18, 22, 19, 17, 21, 24],
            'tiempo_procesamiento': [3.2, 2.8, 2.5, 2.3, 2.1, 2.5, 2.2, 2.0, 2.1, 2.3, 2.0, 1.8]
        }
    }
    return render(request, 'pages/reportes/metricas.html', context)

def actas_configuracion(request):
    """Vista demo para configuración del sistema"""
    context = {
        'parent': 'administracion',
        'segment': 'configuracion',
        'page_title': 'Configuración del Sistema',
        'configuracion_ia': {
            'modelo_transcripcion': 'Whisper v3.0',
            'idioma_principal': 'Español (Ecuador)',
            'nivel_confianza_minimo': 85.0,
            'auto_publicacion': False
        },
        'configuracion_municipal': {
            'nombre_municipio': 'Municipio de Pastaza',
            'direccion': 'Av. Principal 123, Puyo',
            'telefono': '+593-3-2885-123',
            'email': 'municipio@puyo.gob.ec'
        }
    }
    return render(request, 'pages/administracion/configuracion.html', context)

# Portal Ciudadano Views
def portal_ciudadano(request):
  """Vista principal del portal ciudadano para consultar actas"""

  # Parámetros de búsqueda y filtrado
  search_query = request.GET.get('search', '').strip()
  tipo_sesion = request.GET.get('tipo_sesion', '')
  estado = request.GET.get('estado', '')
  acceso = request.GET.get('acceso', '')
  fecha_desde = request.GET.get('fecha_desde', '')
  fecha_hasta = request.GET.get('fecha_hasta', '')
  orden = request.GET.get('orden', '-fecha_sesion')
  page = request.GET.get('page', 1)

  # Base queryset - solo actas que el usuario puede ver
  actas = ActaMunicipal.objects.filter(activo=True)

  # Filtrar por permisos de usuario
  if not request.user.is_authenticated:
    actas = actas.filter(acceso='publico')
  elif not (request.user.is_superuser or request.user.is_staff):
    actas = actas.filter(
      Q(acceso='publico') |
      Q(acceso='restringido')  # Los usuarios autenticados pueden ver restringidas
    )

  # Aplicar filtros de búsqueda
  if search_query:
    from django.db import connection
    
    # Para PostgreSQL, usar SQL con unaccent
    if connection.vendor == 'postgresql':
      with connection.cursor() as cursor:
        # Preparar la consulta con unaccent
        placeholders = ', '.join(['%s'] * 6)  # 6 campos de búsqueda
        sql_condition = f"""
          id IN (
            SELECT id FROM pages_actamunicipal 
            WHERE activo = true AND (
              unaccent(titulo) ILIKE unaccent(%s) OR
              unaccent(numero_acta) ILIKE unaccent(%s) OR  
              unaccent(resumen) ILIKE unaccent(%s) OR
              unaccent(contenido) ILIKE unaccent(%s) OR
              unaccent(palabras_clave) ILIKE unaccent(%s) OR
              unaccent(presidente) ILIKE unaccent(%s)
            )
          )
        """
        search_pattern = f'%{search_query}%'
        actas = actas.extra(where=[sql_condition], params=[search_pattern] * 6)
    else:
      # Fallback para otros backends
      campos_busqueda = ['titulo', 'numero_acta', 'resumen', 'contenido', 'palabras_clave', 'presidente']
      filtros_busqueda = crear_filtros_busqueda_multiple(campos_busqueda, search_query)
      actas = actas.filter(filtros_busqueda)

  if tipo_sesion:
    actas = actas.filter(tipo_sesion__nombre=tipo_sesion)

  if estado:
    actas = actas.filter(estado__nombre=estado)

  if acceso:
    actas = actas.filter(acceso=acceso)

  if fecha_desde:
    try:
      fecha_desde_dt = datetime.strptime(fecha_desde, '%Y-%m-%d')
      actas = actas.filter(fecha_sesion__gte=fecha_desde_dt)
    except ValueError:
      pass

  if fecha_hasta:
    try:
      fecha_hasta_dt = datetime.strptime(fecha_hasta, '%Y-%m-%d')
      actas = actas.filter(fecha_sesion__lte=fecha_hasta_dt)
    except ValueError:
      pass

  # Ordenar resultados (extensión de opciones)
  valid_orders = [
    'fecha_sesion', '-fecha_sesion',
    'fecha_publicacion', '-fecha_publicacion',
    'fecha_creacion', '-fecha_creacion',
    'titulo', '-titulo',
    'numero_acta', '-numero_acta',
    'tipo_sesion__nombre', '-tipo_sesion__nombre',
    'prioridad', '-prioridad',
    'acceso', '-acceso'
  ]
  if orden not in valid_orders:
    orden = '-fecha_sesion'
  actas = actas.order_by(orden)

  # Paginación
  paginator = Paginator(actas, 10)  # 10 actas por página
  actas_page = paginator.get_page(page)

  # Obtener datos para filtros
  tipos_sesion = TipoSesion.objects.filter(activo=True)
  estados = EstadoActa.objects.filter(activo=True)

  # Estadísticas para mostrar
  total_actas = actas.count()
  actas_publicas = actas.filter(acceso='publico').count()
  actas_este_mes = actas.filter(
    fecha_sesion__gte=datetime.now().replace(day=1)
  ).count()

  # Calcular estadísticas adicionales para la cuarta métrica
  from django.db.models import Avg, Count  # noqa: F401 (Count posible uso futuro)
  from datetime import timedelta  # noqa: F401 (posible uso futuro)

  # Días desde la última actualización
  ultima_acta = ActaMunicipal.objects.filter(activo=True).order_by('-fecha_creacion').first()
  dias_ultima_actualizacion = 0
  if ultima_acta:
    dias_ultima_actualizacion = (datetime.now().date() - ultima_acta.fecha_creacion.date()).days

  # Total de visualizaciones este mes
  visualizaciones_mes = VisualizacionActa.objects.filter(
    fecha_visualizacion__gte=datetime.now().replace(day=1)
  ).count()

  # Total de descargas este mes
  descargas_mes = DescargaActa.objects.filter(
    fecha_descarga__gte=datetime.now().replace(day=1)
  ).count()

  # Porcentaje de actas con IA procesada
  actas_con_ia = actas.filter(transcripcion_ia__isnull=False).count()
  porcentaje_ia = round((actas_con_ia / total_actas * 100) if total_actas > 0 else 0, 1)

  # Promedio de precisión IA
  precision_promedio = actas.filter(
    precision_ia__isnull=False
  ).aggregate(Avg('precision_ia'))['precision_ia__avg']
  precision_promedio = round(precision_promedio if precision_promedio else 0, 1)

  context = {
    'actas': actas_page,
    'total_actas': total_actas,
    'actas_publicas': actas_publicas,
    'actas_este_mes': actas_este_mes,
    'visualizaciones_mes': visualizaciones_mes,
    'descargas_mes': descargas_mes,
    'porcentaje_ia': porcentaje_ia,
    'precision_promedio': precision_promedio,
    'dias_ultima_actualizacion': dias_ultima_actualizacion,
    'tipos_sesion': tipos_sesion,
    'estados': estados,
    'search_query': search_query,
    'filters': {
      'tipo_sesion': tipo_sesion,
      'estado': estado,
      'acceso': acceso,
      'fecha_desde': fecha_desde,
      'fecha_hasta': fecha_hasta,
      'orden': orden,
    },
    'acceso_choices': ActaMunicipal.ACCESO_CHOICES,
    # Opciones para select existente (subset principal)
    'orden_choices': [
      ('-fecha_sesion', 'Sesión más recientes'),
      ('fecha_sesion', 'Sesión más antiguos'),
      ('-fecha_publicacion', 'Publicación más recientes'),
      ('fecha_publicacion', 'Publicación más antiguos'),
      ('-fecha_creacion', 'Creación más recientes'),
      ('fecha_creacion', 'Creación más antiguos'),
      ('titulo', 'Título A-Z'),
      ('-titulo', 'Título Z-A'),
      ('numero_acta', 'Número de acta'),
      ('tipo_sesion__nombre', 'Tipo de sesión A-Z'),
      ('-tipo_sesion__nombre', 'Tipo de sesión Z-A'),
      ('prioridad', 'Prioridad (Baja→Alta)'),
      ('-prioridad', 'Prioridad (Alta→Baja)'),
    ],
    # Opciones para toolbar rápida (puede ser subset o igual al completo)
    'orden_toolbar': [
      ('-fecha_sesion', 'Sesión recientes'),
      ('fecha_sesion', 'Sesión antiguos'),
      ('-fecha_publicacion', 'Publicación recientes'),
      ('-fecha_creacion', 'Creación recientes'),
      ('titulo', 'Título A-Z'),
      ('-titulo', 'Título Z-A'),
      ('numero_acta', 'Nº Acta'),
      ('tipo_sesion__nombre', 'Tipo A-Z'),
      ('prioridad', 'Prioridad Baja→Alta'),
      ('-prioridad', 'Prioridad Alta→Baja'),
    ],
  }

  # Etiqueta legible para orden actual
  orden_lookup = dict(context['orden_choices'])
  orden_label = orden_lookup.get(orden, 'Sesión más recientes')
  context['orden_actual_label'] = orden_label
  context['orden_actual'] = orden

  return render(request, 'pages/portal_ciudadano/index.html', context)

def acta_detail(request, pk):
    """Vista detallada de un acta específica"""
    acta = get_object_or_404(ActaMunicipal, pk=pk, activo=True)
    
    # Verificar permisos de acceso
    if not acta.puede_ver_usuario(request.user):
        raise Http404("No tienes permisos para ver esta acta")
    
    # Registrar visualización
    if request.user.is_authenticated:
        VisualizacionActa.objects.create(
            acta=acta,
            usuario=request.user,
            ip_address=get_client_ip(request),
            user_agent=request.META.get('HTTP_USER_AGENT', '')
        )
    else:
        VisualizacionActa.objects.create(
            acta=acta,
            ip_address=get_client_ip(request),
            user_agent=request.META.get('HTTP_USER_AGENT', '')
        )
    
    # Obtener actas relacionadas (mismo tipo de sesión, fechas cercanas)
    actas_relacionadas = ActaMunicipal.objects.filter(
        tipo_sesion=acta.tipo_sesion,
        activo=True
    ).exclude(pk=acta.pk)
    
    if not request.user.is_authenticated:
        actas_relacionadas = actas_relacionadas.filter(acceso='publico')
    elif not (request.user.is_superuser or request.user.is_staff):
        actas_relacionadas = actas_relacionadas.filter(
            Q(acceso='publico') | Q(acceso='restringido')
        )
    
    actas_relacionadas = actas_relacionadas.order_by('-fecha_sesion')[:5]
    
    # Limpiar contenido HTML para mostrar en la página
    try:
        import sys
        import os
        sys.path.append(os.path.join(os.path.dirname(__file__), '../../'))
        from gestion_actas.utils_contenido import limpiar_contenido_html
        contenido_limpio = limpiar_contenido_html(acta.contenido or "")
    except ImportError:
        # Fallback si no se puede importar
        contenido_limpio = acta.contenido or ""
    
    context = {
        'acta': acta,
        'contenido_limpio': contenido_limpio,
        'actas_relacionadas': actas_relacionadas,
        'puede_descargar': acta.archivo_pdf and acta.puede_ver_usuario(request.user),
    }
    
    return render(request, 'pages/portal_ciudadano/detail.html', context)


def convertir_word_a_pdf(ruta_word, nombre_acta):
    """
    Convierte un archivo Word a PDF en formato A4 vertical
    Usa múltiples métodos para asegurar compatibilidad
    """
    try:
        # Generar nombre de archivo PDF
        directorio_word = os.path.dirname(ruta_word)
        nombre_pdf = os.path.splitext(os.path.basename(ruta_word))[0] + '.pdf'
        ruta_pdf = os.path.join(directorio_word, nombre_pdf)
        
        # Método 1: Intentar con pdfkit (requiere wkhtmltopdf)
        try:
            import pdfkit
            from docx import Document
            from docx.shared import Inches
            
            # Leer documento Word y convertir a HTML
            doc = Document(ruta_word)
            html_content = f"""<!DOCTYPE html>
<html lang="es">
<head>
    <meta charset="UTF-8">
    <style>
        body {{
            font-family: 'Times New Roman', serif;
            line-height: 1.6;
            margin: 2cm;
            color: #333;
        }}
        h1, h2, h3 {{ color: #2c3e50; }}
        .center {{ text-align: center; }}
        table {{ width: 100%; border-collapse: collapse; margin: 15px 0; }}
        td {{ padding: 8px; border: 1px solid #ddd; }}
    </style>
</head>
<body>"""
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    if paragraph.style.name.startswith('Heading'):
                        level = '1' if 'Title' in paragraph.style.name else ('2' if paragraph.style.name == 'Heading 1' else '3')
                        html_content += f"<h{level}>{paragraph.text}</h{level}>"
                    else:
                        html_content += f"<p>{paragraph.text}</p>"
            
            # Procesar tablas
            for table in doc.tables:
                html_content += "<table>"
                for row in table.rows:
                    html_content += "<tr>"
                    for cell in row.cells:
                        html_content += f"<td>{cell.text}</td>"
                    html_content += "</tr>"
                html_content += "</table>"
            
            html_content += "</body></html>"
            
            # Opciones PDF A4 vertical
            options = {
                'page-size': 'A4',
                'orientation': 'Portrait',
                'margin-top': '0.75in',
                'margin-right': '0.75in',  
                'margin-bottom': '0.75in',
                'margin-left': '0.75in',
                'encoding': 'UTF-8',
                'no-outline': None,
                'enable-local-file-access': None,
                'print-media-type': None
            }
            
            pdfkit.from_string(html_content, ruta_pdf, options=options)
            
            if os.path.exists(ruta_pdf):
                logger.info(f"PDF generado exitosamente con pdfkit: {ruta_pdf}")
                return ruta_pdf
                
        except ImportError:
            logger.info("pdfkit no disponible, intentando método alternativo")
        except Exception as e:
            logger.warning(f"Error con pdfkit: {str(e)}, intentando método alternativo")
        
        # Método 2: Fallback usando python-docx + reportlab
        try:
            from docx import Document
            from reportlab.lib.pagesizes import A4
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import inch
            from reportlab.lib import colors
            
            # Leer documento Word
            doc = Document(ruta_word)
            
            # Crear documento PDF
            pdf_doc = SimpleDocTemplate(
                ruta_pdf, 
                pagesize=A4,
                rightMargin=0.75*inch,
                leftMargin=0.75*inch,
                topMargin=0.75*inch,
                bottomMargin=0.75*inch
            )
            
            # Estilos
            styles = getSampleStyleSheet()
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Title'],
                fontSize=16,
                spaceAfter=20,
                alignment=1  # Centro
            )
            
            heading_style = ParagraphStyle(
                'CustomHeading',
                parent=styles['Heading1'],
                fontSize=14,
                spaceAfter=12,
                textColor=colors.HexColor('#2c3e50')
            )
            
            normal_style = ParagraphStyle(
                'CustomNormal',
                parent=styles['Normal'],
                fontSize=11,
                spaceAfter=6
            )
            
            story = []
            
            # Procesar párrafos del Word
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    if paragraph.style.name == 'Title':
                        story.append(Paragraph(paragraph.text, title_style))
                    elif paragraph.style.name.startswith('Heading'):
                        story.append(Paragraph(paragraph.text, heading_style))
                    else:
                        story.append(Paragraph(paragraph.text, normal_style))
            
            # Procesar tablas del Word
            for table in doc.tables:
                data = []
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        row_data.append(cell.text)
                    data.append(row_data)
                
                if data:
                    pdf_table = Table(data)
                    pdf_table.setStyle(TableStyle([
                        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#f8f9fa')),
                        ('TEXTCOLOR', (0, 0), (-1, 0), colors.HexColor('#2c3e50')),
                        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                        ('FONTSIZE', (0, 0), (-1, 0), 10),
                        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                        ('BACKGROUND', (0, 1), (-1, -1), colors.white),
                        ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#dee2e6'))
                    ]))
                    story.append(pdf_table)
                    story.append(Spacer(1, 12))
            
            # Construir PDF
            pdf_doc.build(story)
            
            if os.path.exists(ruta_pdf):
                logger.info(f"PDF generado exitosamente con reportlab: {ruta_pdf}")
                return ruta_pdf
                
        except ImportError:
            logger.error("reportlab no disponible")
        except Exception as e:
            logger.error(f"Error con reportlab: {str(e)}")
        
        return None
        
    except Exception as e:
        logger.error(f"Error general convirtiendo Word a PDF: {str(e)}")
        return None


def acta_pdf_view(request, pk):
    """Vista para mostrar el PDF en el navegador (usando exactamente el mismo proceso que Word)"""
    acta = get_object_or_404(ActaMunicipal, pk=pk, activo=True)
    
    # Verificar permisos
    if not acta.puede_ver_usuario(request.user):
        raise Http404("No tienes permisos para ver esta acta")
    
    try:
        # PASO 1: Usar archivo PDF ya generado si existe
        if acta.archivo_pdf and os.path.exists(acta.archivo_pdf.path):
            with open(acta.archivo_pdf.path, 'rb') as pdf_file:
                response = HttpResponse(pdf_file.read(), content_type='application/pdf')
                response['Content-Disposition'] = f'inline; filename="{acta.numero_acta}.pdf"'
                return response
        
        # PASO 2: Usar EXACTAMENTE el mismo proceso que acta_word_download → PDF
        logger.info(f"PDF no encontrado para acta {pk}, generando usando proceso idéntico a Word")
        
        from gestion_actas.generador_documentos import generar_documentos_acta_mejorados
        
        documentos = generar_documentos_acta_mejorados(acta)
        
        # Si se generó Word, convertir a PDF (manteniendo formato idéntico)
        if 'word' in documentos:
            word_path = documentos['word']['ruta']
            pdf_path = convertir_word_a_pdf(word_path, acta.numero_acta)
            
            if pdf_path and os.path.exists(pdf_path):
                # Guardar PDF en modelo para uso futuro
                with open(pdf_path, 'rb') as pdf_file:
                    from django.core.files.base import ContentFile
                    acta.archivo_pdf.save(
                        os.path.basename(pdf_path),
                        ContentFile(pdf_file.read()),
                        save=True
                    )
                
                with open(pdf_path, 'rb') as pdf_file:
                    response = HttpResponse(pdf_file.read(), content_type='application/pdf')
                    response['Content-Disposition'] = f'inline; filename="{acta.numero_acta}.pdf"'
                    return response
        
        raise Http404("No se pudo generar PDF usando proceso de Word")
        
    except Exception as e:
        logger.error(f"Error mostrando PDF del acta {pk}: {str(e)}")
        raise Http404("Error al generar el archivo PDF")

def acta_pdf_download(request, pk):
    """Vista para descargar el PDF (generado desde Word)"""
    acta = get_object_or_404(ActaMunicipal, pk=pk, activo=True)
    
    # Verificar permisos
    if not acta.puede_ver_usuario(request.user):
        raise Http404("No tienes permisos para descargar esta acta")
    
    # Registrar descarga
    if request.user.is_authenticated:
        DescargaActa.objects.create(
            acta=acta,
            usuario=request.user,
            ip_address=get_client_ip(request),
            formato='pdf'
        )
    
    try:
        # Generar documentos (Word primero, luego PDF desde Word)
        from gestion_actas.generador_documentos import generar_documentos_acta_mejorados
        
        documentos = generar_documentos_acta_mejorados(acta)
        
        # Si se generó Word, convertir a PDF
        if 'word' in documentos:
            word_path = documentos['word']['ruta']
            pdf_path = convertir_word_a_pdf(word_path, acta.numero_acta)
            
            if pdf_path and os.path.exists(pdf_path):
                with open(pdf_path, 'rb') as pdf_file:
                    response = HttpResponse(pdf_file.read(), content_type='application/pdf')
                    response['Content-Disposition'] = f'attachment; filename="{acta.numero_acta}.pdf"'
                    return response
        
        # Fallback: usar PDF generado por el sistema original
        if 'pdf' in documentos:
            pdf_path = documentos['pdf']['ruta']
            if os.path.exists(pdf_path):
                with open(pdf_path, 'rb') as pdf_file:
                    response = HttpResponse(pdf_file.read(), content_type='application/pdf')
                    response['Content-Disposition'] = f'attachment; filename="{acta.numero_acta}.pdf"'
                    return response
        
        raise Http404("No se pudo generar el archivo PDF")
        
    except Exception as e:
        logger.error(f"Error descargando PDF del acta {pk}: {str(e)}")
        raise Http404("Error al generar el archivo PDF")


def acta_txt_download(request, pk):
    """Vista para descargar el TXT (usa archivo ya generado durante la publicación)"""
    acta = get_object_or_404(ActaMunicipal, pk=pk, activo=True)
    
    # Verificar permisos
    if not acta.puede_ver_usuario(request.user):
        raise Http404("No tienes permisos para descargar esta acta")
    
    # Registrar descarga
    if request.user.is_authenticated:
        DescargaActa.objects.create(
            acta=acta,
            usuario=request.user,
            ip_address=get_client_ip(request),
            formato='txt'
        )
    else:
        DescargaActa.objects.create(
            acta=acta,
            ip_address=get_client_ip(request),
            formato='txt'
        )
    
    try:
        # Usar archivo TXT ya generado si existe
        if acta.archivo_txt and os.path.exists(acta.archivo_txt.path):
            with open(acta.archivo_txt.path, 'r', encoding='utf-8') as txt_file:
                response = HttpResponse(txt_file.read(), content_type='text/plain; charset=utf-8')
                response['Content-Disposition'] = f'attachment; filename="{acta.numero_acta}.txt"'
                return response
        
        # Fallback: generar TXT al vuelo si no existe el archivo guardado
        logger.warning(f"Archivo TXT no encontrado para acta {pk}, generando al vuelo")
        from gestion_actas.generador_documentos import generar_documentos_acta_mejorados
        
        documentos = generar_documentos_acta_mejorados(acta)
        
        if 'txt' in documentos:
            txt_path = documentos['txt']['ruta']
            with open(txt_path, 'r', encoding='utf-8') as txt_file:
                response = HttpResponse(txt_file.read(), content_type='text/plain; charset=utf-8')
                response['Content-Disposition'] = f'attachment; filename="{acta.numero_acta}.txt"'
                return response
        
        raise Http404("No se pudo generar el archivo TXT")
        
    except Exception as e:
        logger.error(f"Error descargando TXT del acta {pk}: {str(e)}")
        raise Http404("Error al generar el archivo TXT")

def acta_word_download(request, pk):
    """Vista para descargar el Word (usa archivo ya generado durante la publicación)"""
    acta = get_object_or_404(ActaMunicipal, pk=pk, activo=True)
    
    # Verificar permisos
    if not acta.puede_ver_usuario(request.user):
        raise Http404("No tienes permisos para descargar esta acta")
    
    # Registrar descarga
    if request.user.is_authenticated:
        DescargaActa.objects.create(
            acta=acta,
            usuario=request.user,
            ip_address=get_client_ip(request),
            formato='word'
        )
    else:
        DescargaActa.objects.create(
            acta=acta,
            ip_address=get_client_ip(request),
            formato='word'
        )
    
    try:
        # Usar archivo Word ya generado si existe
        if acta.archivo_word and os.path.exists(acta.archivo_word.path):
            with open(acta.archivo_word.path, 'rb') as word_file:
                response = HttpResponse(
                    word_file.read(), 
                    content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                )
                response['Content-Disposition'] = f'attachment; filename="{acta.numero_acta}.docx"'
                return response
        
        # Fallback: generar Word al vuelo si no existe el archivo guardado
        logger.warning(f"Archivo Word no encontrado para acta {pk}, generando al vuelo")
        from gestion_actas.generador_documentos import generar_documentos_acta_mejorados
        
        documentos = generar_documentos_acta_mejorados(acta)
        
        if 'word' in documentos:
            word_path = documentos['word']['ruta']
            with open(word_path, 'rb') as word_file:
                response = HttpResponse(
                    word_file.read(), 
                    content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                )
                response['Content-Disposition'] = f'attachment; filename="{acta.numero_acta}.docx"'
                return response
        
        raise Http404("No se pudo generar el archivo Word")
        
    except Exception as e:
        logger.error(f"Error descargando Word del acta {pk}: {str(e)}")
        raise Http404("Error al generar el archivo Word")



@csrf_exempt
def portal_ciudadano_api(request):
    """API para búsqueda dinámica en el portal ciudadano"""
    if request.method != 'POST':
        return JsonResponse({'error': 'Método no permitido'}, status=405)
    
    try:
        data = json.loads(request.body)
        search_query = data.get('search', '').strip()
        filters = data.get('filters', {})
        page = data.get('page', 1)
        
        # Base queryset
        actas = ActaMunicipal.objects.filter(activo=True)
        
        # Filtrar por permisos
        if not request.user.is_authenticated:
            actas = actas.filter(acceso='publico')
        elif not (request.user.is_superuser or request.user.is_staff):
            actas = actas.filter(
                Q(acceso='publico') | Q(acceso='restringido')
            )
        
        # Aplicar búsqueda
        if search_query:
            from django.db import connection
            
            # Para PostgreSQL, usar SQL con unaccent
            if connection.vendor == 'postgresql':
                with connection.cursor() as cursor:
                    # Preparar la consulta con unaccent
                    sql_condition = f"""
                      id IN (
                        SELECT id FROM pages_actamunicipal 
                        WHERE activo = true AND (
                          unaccent(titulo) ILIKE unaccent(%s) OR
                          unaccent(numero_acta) ILIKE unaccent(%s) OR  
                          unaccent(resumen) ILIKE unaccent(%s) OR
                          unaccent(palabras_clave) ILIKE unaccent(%s)
                        )
                      )
                    """
                    search_pattern = f'%{search_query}%'
                    actas = actas.extra(where=[sql_condition], params=[search_pattern] * 4)
            else:
                # Fallback para otros backends
                campos_busqueda = ['titulo', 'numero_acta', 'resumen', 'palabras_clave']
                filtros_busqueda = crear_filtros_busqueda_multiple(campos_busqueda, search_query)
                actas = actas.filter(filtros_busqueda)
        
        # Aplicar filtros
        if filters.get('tipo_sesion'):
            actas = actas.filter(tipo_sesion__nombre=filters['tipo_sesion'])
        
        if filters.get('estado'):
            actas = actas.filter(estado__nombre=filters['estado'])
        
        if filters.get('acceso'):
            actas = actas.filter(acceso=filters['acceso'])
        
        # Ordenar y paginar
        actas = actas.order_by('-fecha_sesion')
        paginator = Paginator(actas, 10)
        actas_page = paginator.get_page(page)
        
        # Convertir a JSON
        results = []
        for acta in actas_page:
            results.append({
                'id': acta.id,
                'titulo': acta.titulo,
                'numero_acta': acta.numero_acta,
                'fecha_sesion': acta.fecha_sesion.strftime('%d/%m/%Y'),
                'tipo_sesion': acta.tipo_sesion.get_nombre_display() if acta.tipo_sesion else '',
                'resumen': acta.resumen[:200] + '...' if len(acta.resumen) > 200 else acta.resumen,
                'acceso': acta.get_acceso_display(),
                'icono_acceso': acta.icono_acceso,
                'tiene_pdf': bool(acta.archivo_pdf),
                'url_detail': acta.get_absolute_url(),
                'imagen_preview': acta.imagen_preview.url if acta.imagen_preview else None,
            })
        
        return JsonResponse({
            'success': True,
            'results': results,
            'total': paginator.count,
            'page': actas_page.number,
            'num_pages': paginator.num_pages,
            'has_next': actas_page.has_next(),
            'has_previous': actas_page.has_previous(),
        })
        
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)

def get_client_ip(request):
    """Obtener la IP del cliente"""
    x_forwarded_for = request.META.get('HTTP_X_FORWARDED_FOR')
    if x_forwarded_for:
        ip = x_forwarded_for.split(',')[0]
    else:
        ip = request.META.get('REMOTE_ADDR')
    return ip

# ========================================================================
# VISTAS DE TRANSPARENCIA MUNICIPAL
# ========================================================================

def transparencia_index(request):
    """Vista principal del portal de transparencia"""
    from .models import IndicadorTransparencia, MetricaTransparencia, EstadisticaMunicipal, ProyectoMunicipal
    from django.db.models import Sum, Avg
    from datetime import datetime, timedelta
    
    # Obtener métricas principales
    indicadores = IndicadorTransparencia.objects.filter(activo=True).order_by('categoria', 'orden')
    
    # Obtener métricas recientes (último mes)
    fecha_limite = datetime.now().date() - timedelta(days=30)
    metricas_recientes = MetricaTransparencia.objects.filter(
        fecha__gte=fecha_limite,
        indicador__activo=True
    ).select_related('indicador')
    
    # Estadísticas municipales recientes
    estadisticas = EstadisticaMunicipal.objects.filter(
        fecha__gte=fecha_limite
    ).order_by('categoria', '-fecha')[:12]
    
    # Proyectos activos
    proyectos_activos = ProyectoMunicipal.objects.filter(
        estado__in=['en_ejecucion', 'aprobado']
    ).order_by('-fecha_inicio')[:6]
    
    # Métricas de resumen
    total_actas_publicas = ActaMunicipal.objects.filter(
        acceso='publico',
        estado__nombre='publicada'
    ).count()
    
    total_descargas = DescargaActa.objects.count()
    total_visualizaciones = VisualizacionActa.objects.count()
    
    # Presupuesto total proyectos
    presupuesto_total = ProyectoMunicipal.objects.aggregate(
        total=Sum('presupuesto_total')
    )['total'] or 0
    
    presupuesto_ejecutado = ProyectoMunicipal.objects.aggregate(
        ejecutado=Sum('presupuesto_ejecutado')
    )['ejecutado'] or 0
    
    context = {
        'indicadores': indicadores,
        'estadisticas': estadisticas,
        'proyectos_activos': proyectos_activos,
        'metricas': {
            'actas_publicas': total_actas_publicas,
            'descargas': total_descargas,
            'visualizaciones': total_visualizaciones,
            'presupuesto_total': presupuesto_total,
            'presupuesto_ejecutado': presupuesto_ejecutado,
            'porcentaje_ejecucion': round((presupuesto_ejecutado / presupuesto_total * 100), 2) if presupuesto_total > 0 else 0,
        }
    }
    
    return render(request, 'pages/transparencia/index.html', context)

def transparencia_indicadores(request):
    """Vista de indicadores de transparencia"""
    from .models import IndicadorTransparencia, MetricaTransparencia
    
    indicadores = IndicadorTransparencia.objects.filter(activo=True).order_by('categoria', 'orden')
    
    context = {
        'indicadores': indicadores,
    }
    
    return render(request, 'pages/transparencia/indicadores.html', context)

def transparencia_estadisticas(request):
    """Vista de estadísticas municipales"""
    from .models import EstadisticaMunicipal
    
    estadisticas = EstadisticaMunicipal.objects.all().order_by('categoria', '-fecha')
    
    # Agrupar por categoría
    estadisticas_por_categoria = {}
    for estadistica in estadisticas:
        categoria = estadistica.get_categoria_display()
        if categoria not in estadisticas_por_categoria:
            estadisticas_por_categoria[categoria] = []
        estadisticas_por_categoria[categoria].append(estadistica)
    
    context = {
        'estadisticas_por_categoria': estadisticas_por_categoria,
    }
    
    return render(request, 'pages/transparencia/estadisticas.html', context)

def transparencia_proyectos(request):
    """Vista de proyectos municipales"""
    from .models import ProyectoMunicipal
    from django.db.models import Sum
    
    proyectos = ProyectoMunicipal.objects.all().order_by('-fecha_inicio')
    
    # Estadísticas de proyectos
    stats = {
        'total': proyectos.count(),
        'en_ejecucion': proyectos.filter(estado='en_ejecucion').count(),
        'finalizados': proyectos.filter(estado='finalizado').count(),
        'presupuesto_total': proyectos.aggregate(Sum('presupuesto_total'))['presupuesto_total__sum'] or 0,
        'presupuesto_ejecutado': proyectos.aggregate(Sum('presupuesto_ejecutado'))['presupuesto_ejecutado__sum'] or 0,
    }
    
    # Filtros
    categoria_filter = request.GET.get('categoria')
    estado_filter = request.GET.get('estado')
    
    if categoria_filter:
        proyectos = proyectos.filter(categoria=categoria_filter)
    
    if estado_filter:
        proyectos = proyectos.filter(estado=estado_filter)
    
    # Paginación
    paginator = Paginator(proyectos, 12)
    page_number = request.GET.get('page')
    proyectos_page = paginator.get_page(page_number)
    
    context = {
        'proyectos': proyectos_page,
        'stats': stats,
        'categoria_filter': categoria_filter,
        'estado_filter': estado_filter,
        'categorias': ProyectoMunicipal.CATEGORIA_CHOICES,
        'estados': ProyectoMunicipal.ESTADO_CHOICES,
    }
    
    return render(request, 'pages/transparencia/proyectos.html', context)

# ========================================================================
# APIs PARA GRÁFICOS DE TRANSPARENCIA
# ========================================================================

def api_transparencia_metricas(request):
    """API para obtener métricas de transparencia para gráficos"""
    from .models import MetricaTransparencia, IndicadorTransparencia
    from django.db.models import Avg
    from datetime import datetime, timedelta
    
    try:
        categoria = request.GET.get('categoria', 'todas')
        periodo = request.GET.get('periodo', '12')  # meses
        
        # Calcular fecha límite
        fecha_limite = datetime.now().date() - timedelta(days=int(periodo) * 30)
        
        # Filtrar indicadores
        indicadores_query = IndicadorTransparencia.objects.filter(activo=True)
        if categoria != 'todas':
            indicadores_query = indicadores_query.filter(categoria=categoria)
        
        # Obtener métricas
        metricas = MetricaTransparencia.objects.filter(
            indicador__in=indicadores_query,
            fecha__gte=fecha_limite
        ).select_related('indicador').order_by('fecha')
        
        # Agrupar datos por indicador
        datos = {}
        for metrica in metricas:
            nombre = metrica.indicador.nombre
            if nombre not in datos:
                datos[nombre] = {
                    'name': nombre,
                    'data': [],
                    'color': metrica.indicador.color,
                    'tipo': metrica.indicador.tipo,
                }
            datos[nombre]['data'].append({
                'x': metrica.fecha.strftime('%Y-%m-%d'),
                'y': float(metrica.valor)
            })
        
        return JsonResponse({
            'success': True,
            'series': list(datos.values())
        })
        
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)

def api_transparencia_timeline(request):
    """API para timeline de eventos de transparencia"""
    try:
        # Eventos recientes de actas publicadas
        actas_recientes = ActaMunicipal.objects.filter(
            estado__nombre='publicada',
            fecha_publicacion__isnull=False
        ).order_by('-fecha_publicacion')[:10]
        
        eventos = []
        for acta in actas_recientes:
            eventos.append({
                'fecha': acta.fecha_publicacion.strftime('%Y-%m-%d'),
                'titulo': f'Acta Publicada: {acta.numero_acta}',
                'descripcion': acta.titulo,
                'tipo': 'acta',
                'icono': 'fas fa-file-alt',
                'color': '#28a745'
            })
        
        return JsonResponse({
            'success': True,
            'eventos': eventos
        })
        
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)

def api_transparencia_proyectos_stats(request):
    """API para estadísticas de proyectos"""
    from .models import ProyectoMunicipal
    from django.db.models import Count, Sum
    
    try:
        # Estadísticas por estado
        stats_estado = ProyectoMunicipal.objects.values('estado').annotate(
            count=Count('id'),
            presupuesto=Sum('presupuesto_total')
        )
        
        # Estadísticas por categoría
        stats_categoria = ProyectoMunicipal.objects.values('categoria').annotate(
            count=Count('id'),
            presupuesto=Sum('presupuesto_total')
        )
        
        # Formatear datos para gráficos
        datos_estado = []
        for stat in stats_estado:
            datos_estado.append({
                'name': dict(ProyectoMunicipal.ESTADO_CHOICES)[stat['estado']],
                'value': stat['count'],
                'presupuesto': float(stat['presupuesto'] or 0)
            })
        
        datos_categoria = []
        for stat in stats_categoria:
            datos_categoria.append({
                'name': dict(ProyectoMunicipal.CATEGORIA_CHOICES)[stat['categoria']],
                'value': stat['count'],
                'presupuesto': float(stat['presupuesto'] or 0)
            })
        
        return JsonResponse({
            'success': True,
            'por_estado': datos_estado,
            'por_categoria': datos_categoria
        })
        
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)

# ========================================================================
# VISTAS DE EVENTOS MUNICIPALES
# ========================================================================

def eventos_index(request):
    """Vista principal de eventos - redirecciona al calendario"""
    return redirect('eventos_calendario')

def eventos_calendario(request):
    """Vista del calendario de eventos"""
    from .models import EventoMunicipal
    from datetime import datetime, timedelta
    
    # Obtener eventos según visibilidad del usuario
    eventos_query = EventoMunicipal.objects.filter(activo=True)
    
    if not request.user.is_authenticated:
        # Solo eventos públicos para usuarios no autenticados
        eventos_query = eventos_query.filter(visibilidad='publico')
    elif not (request.user.is_superuser or request.user.is_staff):
        # Para usuarios autenticados no admin: públicos + invitados
        from django.db.models import Q
        eventos_query = eventos_query.filter(
            Q(visibilidad='publico') | 
            Q(asistentes_invitados=request.user) |
            Q(organizador=request.user)
        ).distinct()
    
    # Eventos próximos para la barra lateral
    fecha_limite = datetime.now() + timedelta(days=30)
    eventos_proximos = eventos_query.filter(
        fecha_inicio__gte=datetime.now(),
        fecha_inicio__lte=fecha_limite
    ).order_by('fecha_inicio')[:5]
    
    context = {
        'eventos_proximos': eventos_proximos,
        'puede_crear': request.user.is_authenticated,
    }
    
    return render(request, 'pages/eventos/calendario.html', context)

def eventos_nuevo(request):
    """Vista para crear nuevo evento"""
    if not request.user.is_authenticated:
        return redirect('login')
    
    from .forms import EventoForm, DocumentosEventoForm
    from .models import DocumentoEvento, InvitacionExterna
    from django.contrib import messages
    from django.core.mail import send_mail
    from apps.config_system.smtp_service import enviar_email_evento
    from django.conf import settings
    import uuid
    import os
    
    if request.method == 'POST':
        form = EventoForm(request.POST, request.FILES)
        docs_form = DocumentosEventoForm(request.POST, request.FILES)
        
        if form.is_valid():
            evento = form.save(commit=False)
            evento.organizador = request.user
            evento.save()
            form.save_m2m()
            
            # Procesar documentos múltiples
            if docs_form.is_valid():
                archivos = docs_form.cleaned_data.get('archivos', [])
                if archivos:
                    tipo_documento = docs_form.cleaned_data.get('tipo_documento', 'otro')
                    es_publico = docs_form.cleaned_data.get('es_publico', True)
                    
                    documentos_creados = 0
                    for archivo in archivos:
                        documento = DocumentoEvento(
                            evento=evento,
                            nombre=archivo.name,
                            tipo_documento=tipo_documento,
                            archivo=archivo,
                            es_publico=es_publico,
                            subido_por=request.user
                        )
                        documento.save()
                        documentos_creados += 1
                    
                    if documentos_creados > 0:
                        messages.success(request, f'Se subieron {documentos_creados} documento(s) al evento.')
            
            # Procesar invitaciones externas
            emails_externos = form.cleaned_data.get('invitados_externos', [])
            for email in emails_externos:
                if email:
                    # Crear invitación externa con token único
                    invitacion = InvitacionExterna(
                        evento=evento,
                        email=email,
                        token=str(uuid.uuid4()),
                        enviado_por=request.user
                    )
                    invitacion.save()
                    
                    # Enviar email de invitación usando el servicio SMTP personalizado
                    try:
                        exito, mensaje = enviar_email_evento(
                            destinatario=email,
                            evento=evento,
                            tipo_notificacion='invitacion',
                            usuario_solicitante=request.user
                        )
                        
                        if not exito:
                            messages.warning(request, f'Error enviando invitación a {email}: {mensaje}')
                        # El estado se mantiene como 'enviada' que es el default
                    except Exception as e:
                        messages.warning(request, f'Error enviando invitación a {email}: {str(e)}')
            
            messages.success(request, f'Evento "{evento.titulo}" creado exitosamente.')
            if emails_externos:
                messages.info(request, f'Se enviaron {len(emails_externos)} invitación(es) externa(s).')
            
            return redirect('eventos_detalle', pk=evento.pk)
    else:
        form = EventoForm()
        docs_form = DocumentosEventoForm()
    
    context = {
        'form': form,
        'docs_form': docs_form,
        'titulo': 'Crear Nuevo Evento',
    }
    
    return render(request, 'pages/eventos/formulario.html', context)

def eventos_detalle(request, pk):
    """Vista de detalle de evento"""
    from .models import EventoMunicipal, DocumentoEvento, InvitacionExterna
    
    evento = get_object_or_404(EventoMunicipal, pk=pk, activo=True)
    
    # Verificar permisos de visualización
    if not evento.puede_ver_usuario(request.user):
        raise Http404("Evento no encontrado")
    
    # Obtener documentos del evento
    documentos = DocumentoEvento.objects.filter(evento=evento).order_by('-fecha_subida')
    
    # Obtener invitados externos del evento
    invitados_externos = InvitacionExterna.objects.filter(evento=evento).order_by('-fecha_envio')
    
    # Verificar si el usuario puede confirmar asistencia
    puede_confirmar = (
        request.user.is_authenticated and 
        evento.estado in ['programado', 'confirmado'] and
        (evento.es_publico or request.user in evento.asistentes_invitados.all())
    )
    
    # Verificar si ya confirmó asistencia
    ya_confirmo = False
    if request.user.is_authenticated:
        ya_confirmo = request.user in evento.asistentes_confirmados.all()
    
    context = {
        'evento': evento,
        'documentos': documentos,
        'invitados_externos': invitados_externos,
        'puede_confirmar': puede_confirmar,
        'ya_confirmo': ya_confirmo,
        'puede_editar': evento.puede_editar_usuario(request.user),
        # Información de asistentes
        'asistentes_invitados': evento.asistentes_invitados.all(),
        'asistentes_confirmados': evento.asistentes_confirmados.all(),
        'invitados_externos_confirmados': invitados_externos.filter(estado='confirmada'),
        'total_invitados': evento.asistentes_invitados.count() + invitados_externos.count(),
        'total_confirmados': evento.asistentes_confirmados.count() + invitados_externos.filter(estado='confirmada').count(),
    }
    
    return render(request, 'pages/eventos/detalle.html', context)

def eventos_editar(request, pk):
    """Vista para editar evento"""
    from .models import EventoMunicipal, DocumentoEvento, InvitacionExterna
    from .forms import EventoForm, DocumentosEventoForm
    from django.contrib import messages
    import uuid
    
    evento = get_object_or_404(EventoMunicipal, pk=pk, activo=True)
    
    if not evento.puede_editar_usuario(request.user):
        raise Http404("No tienes permisos para editar este evento")
    
    if request.method == 'POST':
        form = EventoForm(request.POST, request.FILES, instance=evento)
        docs_form = DocumentosEventoForm(request.POST, request.FILES)
        
        if form.is_valid():
            evento_actualizado = form.save()
            
            # Procesar documentos múltiples si se enviaron
            if docs_form.is_valid():
                archivos = docs_form.cleaned_data.get('archivos', [])
                if archivos:
                    tipo_documento = docs_form.cleaned_data.get('tipo_documento', 'otro')
                    es_publico = docs_form.cleaned_data.get('es_publico', True)
                    
                    documentos_creados = 0
                    for archivo in archivos:
                        documento = DocumentoEvento(
                            evento=evento_actualizado,
                            nombre=archivo.name,
                            tipo_documento=tipo_documento,
                            archivo=archivo,
                            es_publico=es_publico,
                            subido_por=request.user
                        )
                        documento.save()
                        documentos_creados += 1
                    
                    if documentos_creados > 0:
                        messages.success(request, f'Se agregaron {documentos_creados} documento(s) al evento.')
            
            messages.success(request, f'Evento "{evento_actualizado.titulo}" actualizado exitosamente.')
            return redirect('eventos_detalle', pk=evento_actualizado.pk)
    else:
        form = EventoForm(instance=evento)
        docs_form = DocumentosEventoForm()
    
    context = {
        'form': form,
        'docs_form': docs_form,
        'evento': evento,
        'titulo': f'Editar Evento: {evento.titulo}',
    }
    
    return render(request, 'pages/eventos/formulario.html', context)

def eventos_confirmar_asistencia(request, pk):
    """Vista para confirmar asistencia a evento"""
    from .models import EventoMunicipal, AsistenciaEvento
    
    if not request.user.is_authenticated:
        return JsonResponse({'error': 'Debes estar autenticado'}, status=401)
    
    evento = get_object_or_404(EventoMunicipal, pk=pk, activo=True)
    
    if request.method == 'POST':
        # Verificar si puede confirmar
        if not (evento.es_publico or request.user in evento.asistentes_invitados.all()):
            return JsonResponse({'error': 'No tienes permiso para confirmar asistencia'}, status=403)
        
        # Verificar capacidad
        if evento.capacidad_maxima > 0 and evento.total_asistentes_confirmados >= evento.capacidad_maxima:
            return JsonResponse({'error': 'Evento lleno'}, status=400)
        
        # Crear o actualizar asistencia
        asistencia, created = AsistenciaEvento.objects.get_or_create(
            evento=evento,
            usuario=request.user,
            defaults={'tipo': 'confirmado'}
        )
        
        if created:
            # Agregar a asistentes confirmados
            evento.asistentes_confirmados.add(request.user)
            mensaje = 'Asistencia confirmada exitosamente'
        else:
            mensaje = 'Ya tenías confirmada la asistencia'
        
        return JsonResponse({
            'success': True,
            'mensaje': mensaje,
            'total_confirmados': evento.total_asistentes_confirmados
        })
    
    return JsonResponse({'error': 'Método no permitido'}, status=405)

def eventos_confirmar_externa(request, pk, token):
    """Vista para confirmación de asistencia externa mediante token"""
    from .models import EventoMunicipal, InvitacionExterna
    from django.contrib import messages
    from django.utils import timezone
    
    evento = get_object_or_404(EventoMunicipal, pk=pk, activo=True)
    invitacion = get_object_or_404(InvitacionExterna, evento=evento, token=token)
    
    if request.method == 'POST':
        nombre = request.POST.get('nombre', '').strip()
        if not nombre:
            messages.error(request, 'Por favor ingresa tu nombre.')
        else:
            # Confirmar asistencia
            invitacion.estado = 'confirmada'
            invitacion.nombre = nombre
            invitacion.fecha_respuesta = timezone.now()
            invitacion.save()
            
            messages.success(request, f'¡Gracias {nombre}! Tu asistencia ha sido confirmada.')
            return redirect('eventos_detalle', pk=evento.pk)
    
    # Si ya está confirmado, mostrar información
    if invitacion.estado == 'confirmada':
        messages.info(request, f'Ya has confirmado tu asistencia como {invitacion.nombre}.')
        return redirect('eventos_detalle', pk=evento.pk)
    
    context = {
        'evento': evento,
        'invitacion': invitacion,
    }
    
    return render(request, 'pages/eventos/confirmar_externa.html', context)

# ========================================================================
# APIs PARA CALENDARIO DE EVENTOS
# ========================================================================

def api_eventos_calendario(request):
    """API para obtener eventos del calendario"""
    from .models import EventoMunicipal
    from datetime import datetime
    import json
    
    try:
        # Obtener parámetros de fecha
        start = request.GET.get('start')
        end = request.GET.get('end')
        
        eventos_query = EventoMunicipal.objects.filter(activo=True)
        
        # Filtrar por fechas si se proporcionan
        if start:
            try:
                start_date = datetime.fromisoformat(start.replace('Z', '+00:00'))
                eventos_query = eventos_query.filter(fecha_inicio__gte=start_date)
            except:
                pass
        
        if end:
            try:
                end_date = datetime.fromisoformat(end.replace('Z', '+00:00'))
                eventos_query = eventos_query.filter(fecha_fin__lte=end_date)
            except:
                pass
        
        # Filtrar por visibilidad según usuario
        if not request.user.is_authenticated:
            eventos_query = eventos_query.filter(visibilidad='publico')
        elif not (request.user.is_superuser or request.user.is_staff):
            from django.db.models import Q
            eventos_query = eventos_query.filter(
                Q(visibilidad='publico') | 
                Q(asistentes_invitados=request.user) |
                Q(organizador=request.user)
            ).distinct()
        
        # Convertir a formato FullCalendar
        eventos_calendar = []
        for evento in eventos_query:
            eventos_calendar.append({
                'id': evento.id,
                'title': evento.titulo,
                'start': evento.fecha_inicio.isoformat(),
                'end': evento.fecha_fin.isoformat(),
                'backgroundColor': evento.color_tipo,
                'borderColor': evento.color_tipo,
                'textColor': '#ffffff',
                'url': f'/eventos/{evento.id}/',
                'extendedProps': {
                    'tipo': evento.get_tipo_display(),
                    'estado': evento.get_estado_display(),
                    'ubicacion': evento.ubicacion,
                    'organizador': evento.organizador.get_full_name(),
                    'asistentes': evento.total_asistentes_confirmados,
                    'icono': evento.icono_tipo,
                }
            })
        
        return JsonResponse(eventos_calendar, safe=False)
        
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)

def api_eventos_hoy(request):
    """API para obtener eventos de hoy"""
    from .models import EventoMunicipal
    from datetime import datetime, date
    
    try:
        hoy = date.today()
        
        eventos_query = EventoMunicipal.objects.filter(
            activo=True,
            fecha_inicio__date=hoy
        )
        
        # Filtrar por visibilidad
        if not request.user.is_authenticated:
            eventos_query = eventos_query.filter(visibilidad='publico')
        elif not (request.user.is_superuser or request.user.is_staff):
            from django.db.models import Q
            eventos_query = eventos_query.filter(
                Q(visibilidad='publico') | 
                Q(asistentes_invitados=request.user) |
                Q(organizador=request.user)
            ).distinct()
        
        eventos = []
        for evento in eventos_query.order_by('fecha_inicio'):
            eventos.append({
                'id': evento.id,
                'titulo': evento.titulo,
                'hora': evento.fecha_inicio.strftime('%H:%M'),
                'tipo': evento.get_tipo_display(),
                'ubicacion': evento.ubicacion,
                'color': evento.color_tipo,
                'icono': evento.icono_tipo,
                'url': f'/eventos/{evento.id}/',
            })
        
        return JsonResponse({
            'success': True,
            'fecha': hoy.strftime('%d/%m/%Y'),
            'eventos': eventos,
            'total': len(eventos)
        })
        
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)


@require_http_methods(["POST"])
@csrf_exempt
def reset_acta_publicada(request, pk):
    """
    RESET COMPLETO: Elimina TODA la información procesada del acta,
    dejándola completamente virgen como recién llegada desde generación de actas.
    Solo disponible para administradores.
    """
    import logging
    import os
    logger = logging.getLogger(__name__)
    
    # Verificar permisos de administrador
    logger.info(f"Usuario intentando reset COMPLETO: {request.user} (autenticado: {request.user.is_authenticated}, superuser: {getattr(request.user, 'is_superuser', False)}, staff: {getattr(request.user, 'is_staff', False)})")
    
    if not request.user.is_authenticated:
        messages.error(request, "❌ Debes iniciar sesión para realizar esta acción")
        return redirect('acta_detail', pk=pk)
    
    if not (request.user.is_superuser or request.user.is_staff):
        messages.error(request, "❌ No tienes permisos de administrador para realizar esta acción")
        return redirect('acta_detail', pk=pk)
    
    acta = get_object_or_404(ActaMunicipal, pk=pk, activo=True)
    numero_acta_backup = acta.numero_acta  # Backup para mensajes
    
    try:
        logger.info(f"🔥 INICIANDO RESET COMPLETO del acta {acta.numero_acta} por usuario {request.user.username}")
        
        # ═══════════════════════════════════════════════════════════
        # PASO 1: ELIMINACIÓN TOTAL DE ARCHIVOS FÍSICOS
        # ═══════════════════════════════════════════════════════════
        archivos_eliminados = []
        
        # Eliminar TODOS los archivos relacionados al acta
        archivos_a_eliminar = [
            ('PDF', acta.archivo_pdf),
            ('Word', acta.archivo_word), 
            ('TXT', acta.archivo_txt)
        ]
        
        for tipo, archivo in archivos_a_eliminar:
            if archivo:
                try:
                    if hasattr(archivo, 'path') and os.path.exists(archivo.path):
                        os.remove(archivo.path)
                        archivos_eliminados.append(tipo)
                        logger.info(f"🗑️ Archivo {tipo} eliminado: {archivo.path}")
                    archivo.delete(save=False)
                except Exception as e:
                    logger.error(f"❌ Error eliminando archivo {tipo}: {str(e)}")
        
        # ═══════════════════════════════════════════════════════════
        # PASO 2: LIMPIEZA COMPLETA DEL ACTA EN PORTAL CIUDADANO
        # ═══════════════════════════════════════════════════════════
        logger.info(f"🧹 Limpiando COMPLETAMENTE el acta del portal ciudadano")
        
        # Resetear TODOS los campos procesados a estado virgen (usando campos reales del modelo)
        acta.contenido = ""           # Limpiar contenido procesado
        acta.resumen = ""            # Limpiar resumen
        acta.orden_del_dia = ""      # Limpiar orden del día
        acta.acuerdos = ""           # Limpiar acuerdos
        acta.palabras_clave = ""     # Limpiar palabras clave
        acta.observaciones = ""      # Limpiar observaciones
        
        # Resetear campos de IA
        acta.transcripcion_ia = False        # Marcar como no procesada por IA
        acta.precision_ia = None             # Limpiar métricas IA
        acta.tiempo_procesamiento = None     # Limpiar tiempos
        
        # Resetear fechas de proceso
        acta.fecha_publicacion = None
        
        # Limpiar participantes adicionales
        acta.asistentes = ""         # Limpiar lista de asistentes
        acta.ausentes = ""           # Limpiar lista de ausentes
        
        # Resetear campos de acceso y estado
        acta.acceso = 'publico'      # Resetear a público por defecto
        acta.activo = False          # Despublicar completamente del portal
        acta.prioridad = 'normal'    # Resetear prioridad a normal
        
        acta.save()
        logger.info(f"✅ Acta portal completamente limpiada")
        
        # ═══════════════════════════════════════════════════════════
        # PASO 3: RESET COMPLETO EN GESTOR DE ACTAS
        # ═══════════════════════════════════════════════════════════
        try:
            from gestion_actas.models import GestionActa, EstadoGestionActa
            
            # Buscar el acta en gestor de actas
            acta_gestion = GestionActa.objects.filter(
                numero_acta=numero_acta_backup
            ).first()
            
            if not acta_gestion:
                # Buscar por conexión al portal (si existe)
                acta_gestion = GestionActa.objects.filter(acta_portal=acta).first()
            
            if acta_gestion:
                logger.info(f"🔄 Reseteando COMPLETAMENTE el acta en gestor de actas")
                
                # Obtener estado inicial (en edición/borrador)
                estado_inicial = EstadoGestionActa.objects.filter(
                    codigo__in=['en_edicion', 'borrador', 'creada']
                ).first()
                
                if not estado_inicial:
                    # Crear estado inicial si no existe
                    estado_inicial, created = EstadoGestionActa.objects.get_or_create(
                        codigo='en_edicion',
                        defaults={
                            'nombre': 'En Edición',
                            'descripcion': 'Acta en proceso de edición',
                            'color': '#6c757d',
                            'activo': True
                        }
                    )
                    if created:
                        logger.info(f"📝 Estado 'en_edicion' creado")
                
                # RESET COMPLETO DE CAMPOS EN GESTOR (usando campos reales del modelo)
                acta_gestion.estado = estado_inicial
                acta_gestion.bloqueada_edicion = False
                acta_gestion.contenido_editado = ""  # Limpiar contenido editado
                acta_gestion.observaciones = "🔥 ACTA RESETEADA COMPLETAMENTE - Volver a procesar desde cero"
                
                # Limpiar TODAS las fechas de proceso existentes
                acta_gestion.fecha_enviada_revision = None
                acta_gestion.fecha_aprobacion_final = None
                acta_gestion.fecha_publicacion = None
                
                # Resetear versión y cambios
                acta_gestion.version = 1
                acta_gestion.cambios_realizados = {}
                
                # DESCONECTAR del portal ciudadano
                acta_gestion.acta_portal = None
                
                # Limpiar usuario editor
                acta_gestion.usuario_editor = None
                
                acta_gestion.save()
                logger.info(f"✅ Acta gestor completamente reseteada al estado inicial")
                
                # ═══════════════════════════════════════════════════════════
                # PASO 4: LIMPIAR REGISTROS DE AUDITORÍA Y PROCESAMIENTO
                # ═══════════════════════════════════════════════════════════
                try:
                    # Limpiar registros de visualización
                    from apps.pages.models import VisualizacionActa, DescargaActa
                    VisualizacionActa.objects.filter(acta=acta).delete()
                    DescargaActa.objects.filter(acta=acta).delete()
                    logger.info(f"🧹 Registros de visualización y descarga eliminados")
                    
                    # Limpiar registros de procesamiento de audio (si existen)
                    try:
                        from apps.audio_processing.models import ProcesamientoAudio
                        ProcesamientoAudio.objects.filter(
                            numero_acta=numero_acta_backup
                        ).delete()
                        logger.info(f"🎵 Registros de procesamiento de audio eliminados")
                    except Exception:
                        pass  # Si no existe el modelo, continuar
                    
                except Exception as e:
                    logger.warning(f"⚠️ Error limpiando registros de auditoría: {str(e)}")
                
                # MENSAJE DE ÉXITO COMPLETO
                elementos_reseteados = [
                    "✅ Portal ciudadano (despublicado)",
                    "✅ Gestor de actas (estado inicial)", 
                    "✅ Archivos físicos eliminados",
                    "✅ Contenido y transcripciones",
                    "✅ Fechas de aprobación y publicación",
                    "✅ Metadatos de IA y procesamiento",
                    "✅ Registros de auditoría y descargas"
                ]
                
                messages.success(
                    request, 
                    f"🔥 RESET COMPLETO EXITOSO: {numero_acta_backup}\n\n"
                    f"📋 Elementos completamente eliminados:\n" +
                    "\n".join(elementos_reseteados) +
                    f"\n\n📁 Archivos eliminados: {', '.join(archivos_eliminados) if archivos_eliminados else 'Ninguno'}\n\n"
                    f"🎯 El acta ha vuelto al estado VIRGEN inicial, lista para proceso desde cero."
                )
                
            else:
                logger.warning(f"⚠️ No se encontró el acta {numero_acta_backup} en gestor de actas")
                messages.success(
                    request,
                    f"🔥 RESET COMPLETO DEL PORTAL: {numero_acta_backup}\n\n"
                    f"✅ Acta completamente eliminada del portal ciudadano\n"
                    f"📁 Archivos eliminados: {', '.join(archivos_eliminados) if archivos_eliminados else 'Ninguno'}\n\n"
                    f"⚠️ No se encontró en gestor de actas (puede ser independiente)"
                )
                
        except Exception as e:
            logger.error(f"❌ Error en reset del gestor de actas: {str(e)}")
            messages.error(
                request,
                f"❌ Error al resetear en gestor de actas: {str(e)}\n\n"
                f"✅ Pero el acta fue limpiada del portal ciudadano exitosamente."
            )
        
        # REDIRECT AL PORTAL CIUDADANO (ya no existirá el detalle)
        logger.info(f"🎯 Reset completo finalizado para {numero_acta_backup}")
        return redirect('portal_ciudadano')
        
    except Exception as e:
        logger.error(f"💥 ERROR GENERAL en reset completo del acta {pk}: {str(e)}")
        messages.error(request, f"💥 Error crítico al resetear el acta: {str(e)}")
        return redirect('portal_ciudadano')


@csrf_protect
def reset_sistema_completo_admin(request):
    """
    RESET TOTAL DEL SISTEMA: Resetea TODAS las actas y toda la BD
    al estado de listado en gestión de actas en "edición/depuración"
    Solo disponible para superadministradores
    """
    import logging
    logger = logging.getLogger(__name__)
    
    # Verificar permisos de SUPERADMIN (solo superusuarios)
    if not request.user.is_authenticated or not request.user.is_superuser:
        messages.error(request, "❌ Solo superadministradores pueden ejecutar el reset total del sistema")
        return redirect('index')
    
    if request.method == 'POST':
        try:
            logger.info(f"🔥🔥🔥 INICIANDO RESET TOTAL DEL SISTEMA por {request.user.username}")
            
            # Ejecutar el comando de reset usando call_command
            from django.core.management import call_command
            from io import StringIO
            import sys
            
            # Capturar output del comando
            old_stdout = sys.stdout
            sys.stdout = StringIO()
            
            try:
                # Ejecutar comando de reset completo
                call_command('reset_sistema_completo', confirm=True, verbosity=2)
                comando_output = sys.stdout.getvalue()
                
                logger.info(f"✅ Reset sistema completo ejecutado exitosamente")
                logger.info(f"Output: {comando_output}")
                
                messages.success(
                    request,
                    f"🔥🔥🔥 RESET TOTAL DEL SISTEMA EXITOSO\n\n"
                    f"✅ TODAS las actas reseteadas al estado inicial\n"
                    f"✅ Portal ciudadano completamente limpio\n"
                    f"✅ Gestión de actas en estado 'Edición/Depuración'\n"
                    f"✅ Registros de auditoría y procesamiento limpiados\n"
                    f"✅ Archivos físicos eliminados\n\n"
                    f"🎯 Sistema completamente reseteado - Listo para usar desde cero"
                )
                
            except Exception as e:
                logger.error(f"Error ejecutando comando reset: {str(e)}")
                messages.error(request, f"❌ Error ejecutando reset: {str(e)}")
            
            finally:
                sys.stdout = old_stdout
            
        except Exception as e:
            logger.error(f"💥 ERROR CRÍTICO en reset total del sistema: {str(e)}")
            messages.error(request, f"💥 Error crítico en reset total: {str(e)}")
    
    # Redirect a gestión de actas o página principal
    try:
        return redirect('gestion_actas:listar_actas')
    except:
        return redirect('index')