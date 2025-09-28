"""
Sistema mejorado de generación de documentos para actas municipales
Genera PDF, TXT y Word con formato profesional
"""
import os
import tempfile
from django.conf import settings
from django.template.loader import render_to_string
from datetime import datetime
import logging

logger = logging.getLogger(__name__)

def generar_documentos_acta_mejorados(acta_portal):
    """
    Genera documentos en múltiples formatos con calidad profesional
    
    Args:
        acta_portal: Instancia de ActaMunicipal
        
    Returns:
        dict: Rutas de archivos generados
    """
    try:
        from .utils_contenido import generar_texto_acta_formateado, generar_html_acta_limpio
        
        # Crear directorio para documentos
        fecha_str = datetime.now().strftime('%Y/%m/%d')
        directorio_docs = os.path.join(settings.MEDIA_ROOT, 'actas_publicadas', fecha_str)
        os.makedirs(directorio_docs, exist_ok=True)
        
        # Nombre base para archivos
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        nombre_base = f"{acta_portal.numero_acta}_{timestamp}"
        
        documentos_generados = {}
        
        # 1. Generar TXT
        try:
            texto_formateado = generar_texto_acta_formateado(acta_portal)
            archivo_txt = os.path.join(directorio_docs, f"{nombre_base}.txt")
            
            with open(archivo_txt, 'w', encoding='utf-8') as f:
                f.write(texto_formateado)
            
            documentos_generados['txt'] = {
                'ruta': archivo_txt,
                'nombre': f"{nombre_base}.txt",
                'tipo': 'text/plain',
                'url': f"/media/actas_publicadas/{fecha_str}/{nombre_base}.txt"
            }
            logger.info(f"Archivo TXT generado: {archivo_txt}")
            
        except Exception as e:
            logger.error(f"Error generando TXT: {str(e)}")
        
        # 2. Generar HTML limpio
        try:
            html_limpio = generar_html_acta_limpio(acta_portal)
            
            # Template HTML para documento completo
            html_content = f"""<!DOCTYPE html>
<html lang="es">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{acta_portal.titulo}</title>
    <style>
        body {{
            font-family: 'Times New Roman', serif;
            line-height: 1.6;
            margin: 2cm;
            color: #333;
        }}
        .header {{
            text-align: center;
            border-bottom: 3px solid #2c3e50;
            padding-bottom: 20px;
            margin-bottom: 30px;
        }}
        .header h1 {{
            color: #2c3e50;
            margin: 0;
            font-size: 24px;
        }}
        .header h2 {{
            color: #34495e;
            margin: 5px 0;
            font-size: 18px;
            font-weight: normal;
        }}
        .seccion {{
            margin: 25px 0;
        }}
        .seccion h3 {{
            color: #2c3e50;
            border-bottom: 2px solid #3498db;
            padding-bottom: 5px;
            margin-bottom: 15px;
        }}
        .metadata {{
            background: #f8f9fa;
            padding: 15px;
            border-left: 4px solid #3498db;
            margin: 20px 0;
        }}
        .metadata table {{
            width: 100%;
            border-collapse: collapse;
        }}
        .metadata td {{
            padding: 5px 0;
            border-bottom: 1px solid #dee2e6;
        }}
        .metadata td:first-child {{
            font-weight: bold;
            width: 30%;
        }}
        .footer {{
            margin-top: 50px;
            text-align: center;
            font-size: 12px;
            color: #6c757d;
            border-top: 1px solid #dee2e6;
            padding-top: 20px;
        }}
        .firma-section {{
            margin-top: 50px;
            display: flex;
            justify-content: space-between;
        }}
        .firma {{
            text-align: center;
            width: 45%;
        }}
        .firma-linea {{
            border-bottom: 1px solid #333;
            margin-bottom: 5px;
            height: 50px;
        }}
    </style>
</head>
<body>
    <div class="header">
        <h1>ACTA DE SESIÓN MUNICIPAL</h1>
        <h2>MUNICIPIO DE PASTAZA - ECUADOR</h2>
    </div>
    
    <div class="metadata">
        <table>
            <tr><td>Número de Acta:</td><td>{acta_portal.numero_acta}</td></tr>
            <tr><td>Título:</td><td>{acta_portal.titulo}</td></tr>
            <tr><td>Fecha de Sesión:</td><td>{acta_portal.fecha_sesion.strftime('%d de %B de %Y') if acta_portal.fecha_sesion else 'No disponible'}</td></tr>
            <tr><td>Tipo de Sesión:</td><td>{acta_portal.tipo_sesion.nombre if acta_portal.tipo_sesion else 'No especificado'}</td></tr>
            <tr><td>Presidente:</td><td>{acta_portal.presidente or 'No especificado'}</td></tr>
            <tr><td>Secretario:</td><td>{acta_portal.secretario.get_full_name() if acta_portal.secretario else 'No especificado'}</td></tr>
        </table>
    </div>
    
    <div class="seccion">
        <h3>Resumen Ejecutivo</h3>
        <p>{acta_portal.resumen or 'Sin resumen disponible'}</p>
    </div>
    
    <div class="seccion">
        <h3>Orden del Día</h3>
        {html_limpio['orden_del_dia']}
    </div>
    
    <div class="seccion">
        <h3>Desarrollo de la Sesión</h3>
        {html_limpio['contenido']}
    </div>
    
    <div class="seccion">
        <h3>Acuerdos y Resoluciones</h3>
        {html_limpio['acuerdos']}
    </div>
    
    <div class="seccion">
        <h3>Participantes</h3>
        <p><strong>Asistentes:</strong> {acta_portal.asistentes or 'No especificado'}</p>
    </div>
    
    {f'<div class="seccion"><h3>Observaciones</h3>{html_limpio["observaciones"]}</div>' if acta_portal.observaciones else ''}
    
    <div class="firma-section">
        <div class="firma">
            <div class="firma-linea"></div>
            <p><strong>SECRETARIO MUNICIPAL</strong></p>
        </div>
        <div class="firma">
            <div class="firma-linea"></div>
            <p><strong>ALCALDE MUNICIPAL</strong></p>
        </div>
    </div>
    
    <div class="footer">
        <p><strong>Sistema de Actas Municipales - Municipio de Pastaza</strong></p>
        <p>Documento generado automáticamente el {datetime.now().strftime('%d de %B de %Y a las %H:%M:%S')}</p>
    </div>
</body>
</html>"""
            
            archivo_html = os.path.join(directorio_docs, f"{nombre_base}.html")
            with open(archivo_html, 'w', encoding='utf-8') as f:
                f.write(html_content)
            
            documentos_generados['html'] = {
                'ruta': archivo_html,
                'nombre': f"{nombre_base}.html",
                'tipo': 'text/html',
                'url': f"/media/actas_publicadas/{fecha_str}/{nombre_base}.html"
            }
            logger.info(f"Archivo HTML generado: {archivo_html}")
            
        except Exception as e:
            logger.error(f"Error generando HTML: {str(e)}")
        
        # 3. Generar PDF usando HTML
        try:
            pdf_generado = generar_pdf_desde_html(html_content, nombre_base, directorio_docs)
            if pdf_generado:
                documentos_generados['pdf'] = pdf_generado
                logger.info(f"Archivo PDF generado: {pdf_generado['ruta']}")
            
        except Exception as e:
            logger.error(f"Error generando PDF: {str(e)}")
        
        # 4. Generar Word (DOCX)
        try:
            word_generado = generar_word_documento(acta_portal, nombre_base, directorio_docs)
            if word_generado:
                documentos_generados['word'] = word_generado
                logger.info(f"Archivo Word generado: {word_generado['ruta']}")
            
        except Exception as e:
            logger.error(f"Error generando Word: {str(e)}")
        
        return documentos_generados
        
    except Exception as e:
        logger.error(f"Error general generando documentos: {str(e)}")
        return {}


def generar_pdf_desde_html(html_content, nombre_base, directorio):
    """
    Genera PDF desde contenido HTML usando múltiples opciones de compatibilidad
    """
    try:
        # Opción 1: Intentar con xhtml2pdf (más compatible)
        try:
            from xhtml2pdf import pisa
            
            archivo_pdf = os.path.join(directorio, f"{nombre_base}.pdf")
            
            with open(archivo_pdf, "wb") as pdf_file:
                pisa_status = pisa.CreatePDF(html_content.encode('utf-8'), dest=pdf_file, encoding='utf-8')
                
            if not pisa_status.err:
                fecha_str = datetime.now().strftime('%Y/%m/%d')
                return {
                    'ruta': archivo_pdf,
                    'nombre': f"{nombre_base}.pdf",
                    'tipo': 'application/pdf',
                    'url': f"/media/actas_publicadas/{fecha_str}/{nombre_base}.pdf"
                }
                
        except ImportError:
            logger.info("xhtml2pdf no disponible, probando reportlab")
        except Exception as e:
            logger.error(f"Error con xhtml2pdf: {str(e)}")
        
        # Opción 2: Usar reportlab para generar PDF desde cero
        try:
            pdf_reportlab = generar_pdf_con_reportlab(html_content, nombre_base, directorio)
            if pdf_reportlab:
                return pdf_reportlab
                
        except Exception as e:
            logger.error(f"Error con reportlab: {str(e)}")
        
        # Opción 3: Usar wkhtmltopdf si está disponible
        try:
            import pdfkit
            
            archivo_pdf = os.path.join(directorio, f"{nombre_base}.pdf")
            
            options = {
                'page-size': 'A4',
                'margin-top': '0.75in',
                'margin-right': '0.75in',
                'margin-bottom': '0.75in',
                'margin-left': '0.75in',
                'encoding': "UTF-8",
                'no-outline': None,
                'enable-local-file-access': None
            }
            
            pdfkit.from_string(html_content, archivo_pdf, options=options)
            
            fecha_str = datetime.now().strftime('%Y/%m/%d')
            return {
                'ruta': archivo_pdf,
                'nombre': f"{nombre_base}.pdf",
                'tipo': 'application/pdf',
                'url': f"/media/actas_publicadas/{fecha_str}/{nombre_base}.pdf"
            }
            
        except ImportError:
            logger.info("pdfkit no disponible")
        except Exception as e:
            logger.error(f"Error con pdfkit: {str(e)}")
            
        return None
        
    except Exception as e:
        logger.error(f"Error general generando PDF: {str(e)}")
        return None


def generar_pdf_con_reportlab(html_content, nombre_base, directorio):
    """
    Genera PDF usando reportlab como alternativa compatible
    """
    try:
        from reportlab.lib.pagesizes import A4, letter
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from bs4 import BeautifulSoup
        
        archivo_pdf = os.path.join(directorio, f"{nombre_base}.pdf")
        
        # Crear documento
        doc = SimpleDocTemplate(archivo_pdf, pagesize=A4, 
                              rightMargin=72, leftMargin=72,
                              topMargin=72, bottomMargin=18)
        
        # Estilos
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=18,
            spaceAfter=30,
            alignment=1,  # Centrado
        )
        
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=14,
            spaceAfter=12,
        )
        
        normal_style = styles['Normal']
        
        # Extraer contenido del HTML
        soup = BeautifulSoup(html_content, 'html.parser')
        story = []
        
        # Título
        title = soup.find('h1')
        if title:
            story.append(Paragraph(title.get_text(), title_style))
            story.append(Spacer(1, 12))
        
        # Procesar contenido
        for element in soup.find_all(['h2', 'h3', 'p', 'div']):
            if element.name in ['h2', 'h3']:
                story.append(Paragraph(element.get_text(), heading_style))
            elif element.name in ['p', 'div']:
                text = element.get_text().strip()
                if text:
                    story.append(Paragraph(text, normal_style))
                    story.append(Spacer(1, 6))
        
        # Generar PDF
        doc.build(story)
        
        fecha_str = datetime.now().strftime('%Y/%m/%d')
        return {
            'ruta': archivo_pdf,
            'nombre': f"{nombre_base}.pdf",
            'tipo': 'application/pdf',
            'url': f"/media/actas_publicadas/{fecha_str}/{nombre_base}.pdf"
        }
        
    except ImportError:
        logger.error("reportlab no está instalado")
        return None
    except Exception as e:
        logger.error(f"Error generando PDF con reportlab: {str(e)}")
        return None


def generar_word_documento(acta_portal, nombre_base, directorio):
    """
    Genera documento Word (.docx) con formato profesional
    """
    try:
        from docx import Document
        from docx.shared import Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from .utils_contenido import generar_texto_acta_formateado, html_a_texto_formateado
        
        # Crear documento
        doc = Document()
        
        # Título principal
        title = doc.add_heading('ACTA DE SESIÓN MUNICIPAL', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        subtitle = doc.add_heading('MUNICIPIO DE PASTAZA - ECUADOR', level=1)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Información general
        doc.add_heading('INFORMACIÓN GENERAL', level=2)
        
        info_table = doc.add_table(rows=6, cols=2)
        info_table.style = 'Light Grid Accent 1'
        
        info_data = [
            ('Número de Acta:', acta_portal.numero_acta),
            ('Título:', acta_portal.titulo),
            ('Fecha de Sesión:', acta_portal.fecha_sesion.strftime('%d de %B de %Y') if acta_portal.fecha_sesion else 'No disponible'),
            ('Tipo de Sesión:', acta_portal.tipo_sesion.nombre if acta_portal.tipo_sesion else 'No especificado'),
            ('Presidente:', acta_portal.presidente or 'No especificado'),
            ('Secretario:', acta_portal.secretario.get_full_name() if acta_portal.secretario else 'No especificado'),
        ]
        
        for i, (campo, valor) in enumerate(info_data):
            info_table.cell(i, 0).text = campo
            info_table.cell(i, 1).text = str(valor)
        
        # Resumen
        doc.add_heading('RESUMEN EJECUTIVO', level=2)
        doc.add_paragraph(acta_portal.resumen or 'Sin resumen disponible')
        
        # Orden del día
        if acta_portal.orden_del_dia:
            doc.add_heading('ORDEN DEL DÍA', level=2)
            texto_orden = html_a_texto_formateado(acta_portal.orden_del_dia)
            doc.add_paragraph(texto_orden)
        
        # Contenido
        doc.add_heading('DESARROLLO DE LA SESIÓN', level=2)
        texto_contenido = html_a_texto_formateado(acta_portal.contenido) if acta_portal.contenido else 'Sin contenido disponible'
        doc.add_paragraph(texto_contenido)
        
        # Acuerdos
        if acta_portal.acuerdos:
            doc.add_heading('ACUERDOS Y RESOLUCIONES', level=2)
            texto_acuerdos = html_a_texto_formateado(acta_portal.acuerdos)
            doc.add_paragraph(texto_acuerdos)
        
        # Participantes
        doc.add_heading('PARTICIPANTES', level=2)
        doc.add_paragraph(f"Asistentes: {acta_portal.asistentes or 'No especificado'}")
        
        # Observaciones
        if acta_portal.observaciones:
            doc.add_heading('OBSERVACIONES', level=2)
            texto_observaciones = html_a_texto_formateado(acta_portal.observaciones)
            doc.add_paragraph(texto_observaciones)
        
        # Firmas
        doc.add_page_break()
        doc.add_paragraph()
        doc.add_paragraph()
        
        # Tabla de firmas
        firma_table = doc.add_table(rows=3, cols=2)
        firma_table.cell(0, 0).text = ""
        firma_table.cell(0, 1).text = ""
        firma_table.cell(1, 0).text = "_" * 30
        firma_table.cell(1, 1).text = "_" * 30
        firma_table.cell(2, 0).text = "SECRETARIO MUNICIPAL"
        firma_table.cell(2, 1).text = "ALCALDE MUNICIPAL"
        
        for row in firma_table.rows:
            for cell in row.cells:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Footer
        doc.add_paragraph()
        footer_p = doc.add_paragraph(f"Documento generado automáticamente el {datetime.now().strftime('%d de %B de %Y a las %H:%M:%S')}")
        footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Guardar archivo
        archivo_word = os.path.join(directorio, f"{nombre_base}.docx")
        doc.save(archivo_word)
        
        fecha_str = datetime.now().strftime('%Y/%m/%d')
        return {
            'ruta': archivo_word,
            'nombre': f"{nombre_base}.docx",
            'tipo': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            'url': f"/media/actas_publicadas/{fecha_str}/{nombre_base}.docx"
        }
        
    except ImportError:
        logger.error("python-docx no está instalado")
        return None
    except Exception as e:
        logger.error(f"Error generando documento Word: {str(e)}")
        return None